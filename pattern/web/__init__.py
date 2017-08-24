#### PATTERN | WEB #################################################################################
# -*- coding: utf-8 -*-
# Copyright (c) 2010 University of Antwerp, Belgium
# Author: Tom De Smedt <tom@organisms.be>
# License: BSD (see LICENSE.txt for details).
# http://www.clips.ua.ac.be/pages/pattern

####################################################################################################
# Python API interface for various web services (Google, Twitter, Wikipedia, ...)

from __future__ import absolute_import
from __future__ import unicode_literals
from __future__ import print_function
from __future__ import division

from builtins import str, bytes, dict, int, chr
from builtins import map, filter, zip
from builtins import object, range, next

import os
import sys
import threading
import time
import socket
import ssl

from io import open

try:
    # Python 3
    from urllib.parse import urlparse, urljoin, urlsplit, urlencode, quote_plus, unquote_plus
    from urllib.request import urlopen, Request, HTTPHandler, HTTPRedirectHandler, ProxyHandler, HTTPCookieProcessor, install_opener, build_opener
    from urllib.error import HTTPError as UrllibHTTPError
    from urllib.error import URLError as UrllibURLError
except ImportError:
    # Python 2
    from urlparse import urlparse, urljoin, urlsplit
    from urllib import urlencode, quote_plus, unquote_plus
    from urllib2 import urlopen, Request, HTTPHandler, HTTPRedirectHandler, ProxyHandler, HTTPCookieProcessor, install_opener, build_opener
    from urllib2 import HTTPError as UrllibHTTPError
    from urllib2 import URLError as UrllibURLError

import base64

from html.entities import name2codepoint

try:
    # Python 2
    import httplib
except ImportError:
    # Python 3
    import http.client as httplib

from html.parser import HTMLParser as _HTMLParser

try:
    # Python 3
    import cookielib
except ImportError:
    # Python 2
    import http.cookiejar as cookielib
import re
import xml.dom.minidom
import unicodedata
import string
try:
    # Python 2
    from cStringIO import StringIO
except ImportError:
    # Python 3
    from io import StringIO
import bisect
import itertools
try:
    # Python 2
    import new
except ImportError:
    # Python 3: We don't actually need it (in this case)
    new = None
import feedparser
import json

from . import api
from . import oauth
from . import locale

import bs4 as BeautifulSoup

try:
    # Import persistent Cache.
    # If this module is used separately,
    # a dict is used (i.e. this Python session only).
    from .cache import Cache, cache, TMP
except:
    cache = {}

try:
    from .imap import Mail, MailFolder, Message, GMAIL
    from .imap import MailError, MailServiceError, MailLoginError, MailNotLoggedIn
    from .imap import FROM, SUBJECT, DATE, BODY, ATTACHMENTS
except:
    pass

try:
    MODULE = os.path.dirname(os.path.realpath(__file__))
except:
    MODULE = ""

#### UNICODE #######################################################################################
# Latin-1 (ISO-8859-1) encoding is identical to Windows-1252 except for the code points 128-159:
# Latin-1 assigns control codes in this range, Windows-1252 has characters, punctuation, symbols
# assigned to these code points.

from pattern.helpers import encode_string, decode_string

u = decode_utf8 = decode_string
s = encode_utf8 = encode_string

GREMLINS = set([
    0x0152, 0x0153, 0x0160, 0x0161, 0x0178, 0x017E, 0x017D, 0x0192, 0x02C6,
    0x02DC, 0x2013, 0x2014, 0x201A, 0x201C, 0x201D, 0x201E, 0x2018, 0x2019,
    0x2020, 0x2021, 0x2022, 0x2026, 0x2030, 0x2039, 0x203A, 0x20AC, 0x2122
])


def fix(s, ignore=""):
    """ Returns a Unicode string that fixes common encoding problems (Latin-1, Windows-1252).
        For example: fix("clichÃ©") => "cliché".
    """
    # http://blog.luminoso.com/2012/08/20/fix-unicode-mistakes-with-python/
    if not isinstance(s, str):
        s = s.decode("utf-8")
        # If this doesn't work,
        # copy & paste string in a Unicode .txt,
        # and then pass open(f).read() to fix().
    u = []
    i = 0
    for j, ch in enumerate(s):
        if ch in ignore:
            continue
        if ord(ch) < 128: # ASCII
            continue
        if ord(ch) in GREMLINS:
            ch = ch.encode("windows-1252")
        else:
            try:
                ch = ch.encode("latin-1")
            except:
                ch = ch.encode("utf-8")
        u.append(s[i:j].encode("utf-8"))
        u.append(ch)
        i = j + 1
    u.append(s[i:].encode("utf-8"))
    u = b"".join(u)
    u = u.decode("utf-8", "replace")
    u = u.replace("\n", "\n ")
    u = u.split(" ")
    # Revert words that have the replacement character,
    # i.e., fix("cliché") should not return "clich�".
    for i, (w1, w2) in enumerate(zip(s.split(" "), u)):
        if "\ufffd" in w2: # �
            u[i] = w1
    u = " ".join(u)
    u = u.replace("\n ", "\n")
    return u


def latin(s):
    """ Returns True if the string contains only Latin-1 characters
        (no Chinese, Japanese, Arabic, Cyrillic, Hebrew, Greek, ...).
    """
    if not isinstance(s, str):
        s = s.decode("utf-8")
    return all(unicodedata.name(ch).startswith("LATIN") for ch in s if ch.isalpha())

# For clearer source code:
bytestring = b = s

#### ASYNCHRONOUS REQUEST ##########################################################################


class AsynchronousRequest(object):

    def __init__(self, function, *args, **kwargs):
        """ Executes the function in the background.
            AsynchronousRequest.done is False as long as it is busy, but the program will not halt in the meantime.
            AsynchronousRequest.value contains the function's return value once done.
            AsynchronousRequest.error contains the Exception raised by an erronous function.
            For example, this is useful for running live web requests while keeping an animation running.
            For good reasons, there is no way to interrupt a background process (i.e. Python thread).
            You are responsible for ensuring that the given function doesn't hang.
        """
        self._response = None # The return value of the given function.
        self._error = None # The exception (if any) raised by the function.
        self._time = time.time()
        self._function = function
        self._thread = threading.Thread(target=self._fetch, args=(function,) + args, kwargs=kwargs)
        self._thread.start()

    def _fetch(self, function, *args, **kwargs):
        """ Executes the function and sets AsynchronousRequest.response.
        """
        try:
            self._response = function(*args, **kwargs)
        except Exception as e:
            self._error = e

    def now(self):
        """ Waits for the function to finish and yields its return value.
        """
        self._thread.join()
        return self._response

    @property
    def elapsed(self):
        return time.time() - self._time

    @property
    def done(self):
        return not self._thread.isAlive()

    @property
    def value(self):
        return self._response

    @property
    def error(self):
        return self._error

    def __repr__(self):
        return "AsynchronousRequest(function='%s')" % self._function.__name__


def asynchronous(function, *args, **kwargs):
    """ Returns an AsynchronousRequest object for the given function.
    """
    return AsynchronousRequest(function, *args, **kwargs)

send = asynchronous

#### URL ###########################################################################################

# User agent and referrer.
# Used to identify the application accessing the web.
USER_AGENT = "Pattern/2.6 +http://www.clips.ua.ac.be/pattern"
REFERRER = "http://www.clips.ua.ac.be/pattern"

# Mozilla user agent.
# Websites can include code to block out any application except browsers.
MOZILLA = "Mozilla/5.0"

# HTTP request method.
GET = "get"  # Data is encoded in the URL.
POST = "post" # Data is encoded in the message body.

# URL parts.
# protocol://username:password@domain:port/path/page?query_string#anchor
PROTOCOL, USERNAME, PASSWORD, DOMAIN, PORT, PATH, PAGE, QUERY, ANCHOR = \
    "protocol", "username", "password", "domain", "port", "path", "page", "query", "anchor"

# MIME type.
MIMETYPE_WEBPAGE    = ["text/html"]
MIMETYPE_STYLESHEET = ["text/css"]
MIMETYPE_PLAINTEXT  = ["text/plain"]
MIMETYPE_PDF        = ["application/pdf"]
MIMETYPE_NEWSFEED   = ["application/rss+xml", "application/atom+xml"]
MIMETYPE_IMAGE      = ["image/gif", "image/jpeg", "image/png", "image/tiff"]
MIMETYPE_AUDIO      = ["audio/mpeg", "audio/mp4", "audio/x-aiff", "audio/x-wav"]
MIMETYPE_VIDEO      = ["video/mpeg", "video/mp4", "video/avi", "video/quicktime", "video/x-flv"]
MIMETYPE_ARCHIVE    = ["application/x-stuffit", "application/x-tar", "application/zip"]
MIMETYPE_SCRIPT     = ["application/javascript", "application/ecmascript"]


def extension(filename):
    """ Returns the extension in the given filename: "cat.jpg" => ".jpg".
    """
    return os.path.splitext(filename)[1]


def urldecode(query):
    """ Inverse operation of urlencode.
        Returns a dictionary of (name, value)-items from a URL query string.
    """
    def _format(s):
        if s == "" or s == "None":
            return None
        if s.lstrip("-").isdigit():
            return int(s)
        try:
            return float(s)
        except:
            return s
    if query:
        query = query.lstrip("?").split("&")
        query = ((kv.split("=") + [None])[:2] for kv in query)
        if sys.version > "3":
            query = ((u(unquote_plus(k)), _format(u(unquote_plus(v)))) for k, v in query if k != "")
        else:
            query = ((u(unquote_plus(bytestring(k))), _format(u(unquote_plus(bytestring(v))))) for k, v in query if k != "")
        return dict(query)
    return {}

url_decode = urldecode


def proxy(host, protocol="https"):
    """ Returns the value for the URL.open() proxy parameter.
        - host: host address of the proxy server.
    """
    return (host, protocol)


class Error(Exception):
    """ Base class for pattern.web errors.
    """

    def __init__(self, *args, **kwargs):
        Exception.__init__(self, *args)
        self.src = kwargs.pop("src", None)
        self.url = kwargs.pop("url", None)

    @property
    def headers(self):
        return dict(list(self.src.headers.items()))


class URLError(Error):
    pass # URL contains errors (e.g. a missing t in htp://).


class URLTimeout(URLError):
    pass # URL takes to long to load.


class HTTPError(URLError):
    pass # URL causes an error on the contacted server.


class HTTP301Redirect(HTTPError):
    pass # Too many redirects.
         # The site may be trying to set a cookie and waiting for you to return it,
         # or taking other measures to discern a browser from a script.
         # For specific purposes you should build your own urllib2.HTTPRedirectHandler
         # and pass it to urllib2.build_opener() in URL.open()


class HTTP400BadRequest(HTTPError):
    pass # URL contains an invalid request.


class HTTP401Authentication(HTTPError):
    pass # URL requires a login and password.


class HTTP403Forbidden(HTTPError):
    pass # URL is not accessible (user-agent?)


class HTTP404NotFound(HTTPError):
    pass # URL doesn't exist on the internet.


class HTTP414RequestURITooLong(HTTPError):
    pass # URL is too long.


class HTTP420Error(HTTPError):
    pass # Used by Twitter for rate limiting.


class HTTP429TooMayRequests(HTTPError):
    pass # Used by Twitter for rate limiting.


class HTTP500InternalServerError(HTTPError):
    pass # Generic server error.


class HTTP503ServiceUnavailable(HTTPError):
    pass # Used by Bing for rate limiting.


class URL(object):

    def __init__(self, string="", method=GET, query={}, **kwargs):
        """ URL object with the individual parts available as attributes:
            For protocol://username:password@domain:port/path/page?query_string#anchor:
            - URL.protocol: http, https, ftp, ...
            - URL.username: username for restricted domains.
            - URL.password: password for restricted domains.
            - URL.domain  : the domain name, e.g. nodebox.net.
            - URL.port    : the server port to connect to.
            - URL.path    : the server path of folders, as a list, e.g. ['news', '2010']
            - URL.page    : the page name, e.g. page.html.
            - URL.query   : the query string as a dictionary of (name, value)-items.
            - URL.anchor  : the page anchor.
            If method is POST, the query string is sent with HTTP POST.
        """
        self.__dict__["method"] = method # Use __dict__ directly since __setattr__ is overridden.
        self.__dict__["_string"] = u(string)
        self.__dict__["_parts"] = None
        self.__dict__["_headers"] = None
        self.__dict__["_redirect"] = None
        if isinstance(string, URL):
            self.__dict__["method"] = string.method
            self.query.update(string.query)
        if len(query) > 0:
            # Requires that we parse the string first (see URL.__setattr__).
            self.query.update(query)
        if len(kwargs) > 0:
            # Requires that we parse the string first (see URL.__setattr__).
            self.parts.update(kwargs)

    def _parse(self):
        """ Parses all the parts of the URL string to a dictionary.
            URL format: protocal://username:password@domain:port/path/page?querystring#anchor
            For example: http://user:pass@example.com:992/animal/bird?species=seagull&q#wings
            This is a cached method that is only invoked when necessary, and only once.
        """
        p = urlsplit(self._string)
        P = {PROTOCOL: p[0],            # http
             USERNAME: "",             # user
             PASSWORD: "",             # pass
               DOMAIN: p[1],            # example.com
                 PORT: "",             # 992
                 PATH: p[2],            # [animal]
                 PAGE: "",             # bird
                QUERY: urldecode(p[3]), # {"species": "seagull", "q": None}
               ANCHOR: p[4]             # wings
        }

        # Split the username and password from the domain.
        if "@" in P[DOMAIN]:
            P[USERNAME], \
            P[PASSWORD] = (p[1].split("@")[0].split(":") + [""])[:2]
            P[DOMAIN] = p[1].split("@")[1]
        # Split the port number from the domain.
        if ":" in P[DOMAIN]:
            P[DOMAIN], \
            P[PORT] = P[DOMAIN].split(":")
            P[PORT] = P[PORT].isdigit() and int(P[PORT]) or P[PORT]
        # Split the base page from the path.
        if "/" in P[PATH]:
            P[PAGE] = p[2].split("/")[-1]
            P[PATH] = p[2][:len(p[2]) - len(P[PAGE])].strip("/").split("/")
            P[PATH] = list(filter(lambda v: v != "", P[PATH]))
        else:
            P[PAGE] = p[2].strip("/")
            P[PATH] = []
        self.__dict__["_parts"] = P

    # URL.string yields unicode(URL) by joining the different parts,
    # if the URL parts have been modified.
    def _get_string(self): return str(self)

    def _set_string(self, v):
        self.__dict__["_string"] = u(v)
        self.__dict__["_parts"] = None

    string = property(_get_string, _set_string)

    @property
    def parts(self):
        """ Yields a dictionary with the URL parts.
        """
        if not self._parts:
            self._parse()
        return self._parts

    @property
    def querystring(self):
        """ Yields the URL querystring: "www.example.com?page=1" => "page=1"
        """
        s = self.parts[QUERY].items()
        s = dict((bytestring(k), v if v is not None else "") for k, v in s)
        if sys.version > "3":
            # Python 3
            s = urlencode(s)
        else:
            # Python 2: urlencode() expects byte strings
            t = {key: value.encode("utf-8") if isinstance(value, str) else value for key, value in s.items()}
            s = urlencode(t).decode("utf-8")
        return s

    def __getattr__(self, k):
        if k in self.__dict__:
            return self.__dict__[k]
        if k in self.parts:
            return self.__dict__["_parts"][k]
        raise AttributeError("'URL' object has no attribute '%s'" % k)

    def __setattr__(self, k, v):
        if k in self.__dict__:
            self.__dict__[k] = u(v)
            return
        if k == "string":
            self._set_string(v)
            return
        if k == "query":
            self.parts[k] = v
            return
        if k in self.parts:
            self.__dict__["_parts"][k] = u(v)
            return
        raise AttributeError("'URL' object has no attribute '%s'" % k)

    def open(self, timeout=10, proxy=None, user_agent=USER_AGENT, referrer=REFERRER, authentication=None):
        """ Returns a connection to the url from which data can be retrieved with connection.read().
            When the timeout amount of seconds is exceeded, raises a URLTimeout.
            When an error occurs, raises a URLError (e.g. HTTP404NotFound).
        """
        url = self.string
        # Handle local files directly
        if os.path.exists(url):
            return urlopen(url)
        # Handle method=POST with query string as a separate parameter.
        post = self.method == POST and self.querystring or None
        socket.setdefaulttimeout(timeout)
        # Handle proxies and cookies.
        handlers = []
        if proxy:
            handlers.append(ProxyHandler({proxy[1]: proxy[0]}))
        handlers.append(HTTPCookieProcessor(cookielib.CookieJar()))
        handlers.append(HTTPHandler)
        install_opener(build_opener(*handlers))
        # Send request.
        try:
            request = Request(url, post, {
                        "User-Agent": user_agent,
                           "Referer": referrer
                         })
            # Basic authentication is established with authentication=(username, password).
            if authentication is not None:
                authentication = tuple(encode_utf8(x) for x in authentication)
                request.add_header("Authorization", "Basic %s" %
                    decode_utf8(base64.b64encode(b'%s:%s' % authentication)))
            return urlopen(request)
        except UrllibHTTPError as e:
            if e.code == 301:
                raise HTTP301Redirect(src=e, url=url)
            if e.code == 400:
                raise HTTP400BadRequest(src=e, url=url)
            if e.code == 401:
                raise HTTP401Authentication(src=e, url=url)
            if e.code == 403:
                raise HTTP403Forbidden(src=e, url=url)
            if e.code == 404:
                raise HTTP404NotFound(src=e, url=url)
            if e.code == 414:
                raise HTTP414RequestURITooLong(src=e, url=url)
            if e.code == 420:
                raise HTTP420Error(src=e, url=url)
            if e.code == 429:
                raise HTTP429TooMayRequests(src=e, url=url)
            if e.code == 500:
                raise HTTP500InternalServerError(src=e, url=url)
            if e.code == 503:
                raise HTTP503ServiceUnavailable(src=e, url=url)
            raise HTTPError(str(e), src=e, url=url)
        except httplib.BadStatusLine as e:
            raise HTTPError(str(e), src=e, url=url)
        except socket.timeout as e:
            raise URLTimeout(src=e, url=url)
        except socket.error as e:
            if "timed out" in str((e.args + ("", ""))[0]) \
            or "timed out" in str((e.args + ("", ""))[1]):
                raise URLTimeout(src=e, url=url)
            raise URLError(str(e), src=e, url=url)
        except UrllibURLError as e:
            if "timed out" in str(e.reason):
                raise URLTimeout(src=e, url=url)
            raise URLError(str(e), src=e, url=url)
        except ValueError as e:
            raise URLError(str(e), src=e, url=url)

    def download(self, timeout=10, cached=True, throttle=0, proxy=None, user_agent=USER_AGENT, referrer=REFERRER, authentication=None, unicode=False, **kwargs):
        """ Downloads the content at the given URL (by default it will be cached locally).
            Unless unicode=False, the content is returned as a unicode string.
        """
        # Filter OAuth parameters from cache id (they will be unique for each request).
        if self._parts is None and self.method == GET and "oauth_" not in self._string:
            id = self._string
        else:
            id = repr(self.parts)
            id = re.sub("u{0,1}'oauth_.*?': u{0,1}'.*?', ", "", id)
        # Keep a separate cache of unicode and raw download for same URL.
        if unicode is True:
            id = "u" + id
        if cached and id in cache:
            if isinstance(cache, dict): # Not a Cache object.
                return cache[id]
            if unicode is True:
                return cache[id]
            if unicode is False:
                return cache.get(id, unicode=False)
        t = time.time()
        # Open a connection with the given settings, read it and (by default) cache the data.
        try:
            data = self.open(timeout, proxy, user_agent, referrer, authentication).read()
        except socket.timeout as e:
            raise URLTimeout(src=e, url=self.string)
        if unicode is True:
            data = u(data)
        if cached:
            cache[id] = data
        if throttle:
            time.sleep(max(throttle - (time.time() - t), 0))
        return data

    def read(self, *args, **kwargs):
        return self.open(**kwargs).read(*args)

    @property
    def exists(self, timeout=10):
        """ Yields False if the URL generates a HTTP404NotFound error.
        """
        try:
            self.open(timeout)
        except HTTP404NotFound:
            return False
        except HTTPError:
            return True
        except URLTimeout:
            return True
        except URLError:
            return False
        except:
            return True
        return True

    @property
    def mimetype(self, timeout=10):
        """ Yields the MIME-type of the document at the URL, or None.
            MIME is more reliable than simply checking the document extension.
            You can then do: URL.mimetype in MIMETYPE_IMAGE.
        """
        try:
            return self.headers["content-type"].split(";")[0]
        except KeyError:
            return None

    @property
    def headers(self, timeout=10):
        """ Yields a dictionary with the HTTP response headers.
        """
        if self.__dict__["_headers"] is None:
            try:
                h = dict(self.open(timeout).info())
            except URLError:
                h = {}
            self.__dict__["_headers"] = h

        # Backward compatibility (Python 2)
        if "Content-Type" in self.__dict__["_headers"]:
            self.__dict__["_headers"]["content-type"] = self.__dict__["_headers"]["Content-Type"]

        return self.__dict__["_headers"]

    @property
    def redirect(self, timeout=10):
        """ Yields the redirected URL, or None.
        """
        if self.__dict__["_redirect"] is None:
            try:
                r = u(self.open(timeout).geturl())
            except URLError:
                r = None
            self.__dict__["_redirect"] = r != self.string and r or ""
        return self.__dict__["_redirect"] or None

    def __str__(self):
        # The string representation includes the query attributes with HTTP GET.
        P = self.parts
        u = []
        if P[PROTOCOL]:
            u.append("%s://" % P[PROTOCOL])
        if P[USERNAME]:
            u.append("%s:%s@" % (P[USERNAME], P[PASSWORD]))
        if P[DOMAIN]:
            u.append(P[DOMAIN])
        if P[PORT]:
            u.append(":%s" % P[PORT])
        if P[PORT] or P[DOMAIN] and not P[PATH] and not P[PAGE]:
            u.append("/")
        if P[PATH]:
            u.append("/%s/" % "/".join(P[PATH]))
        if P[PAGE] and len(u) > 0:
            u[-1] = u[-1].rstrip("/")
        if P[PAGE]:
            u.append("/%s" % P[PAGE])
        if P[QUERY] and self.method == GET:
            u.append("?%s" % self.querystring)
        if P[ANCHOR]:
            u.append("#%s" % P[ANCHOR])
        u = "".join(u)
        u = u.lstrip("/")
        return u

    def __repr__(self):
        return "URL(%s, method=%s)" % (repr(self.string), repr(self.method))

    def copy(self):
        return URL(self.string, self.method, self.query)


def download(url="", method=GET, query={}, timeout=10, cached=True, throttle=0, proxy=None, user_agent=USER_AGENT, referrer=REFERRER, authentication=None, unicode=False):
    """ Downloads the content at the given URL (by default it will be cached locally).
        Unless unicode=False, the content is returned as a unicode string.
    """
    return URL(url, method, query).download(timeout, cached, throttle, proxy, user_agent, referrer, authentication, unicode)

#url = URL("http://user:pass@example.com:992/animal/bird?species#wings")
#print(url.parts)
#print(url.query)
#print(url.string)

#--- STREAMING URL BUFFER --------------------------------------------------------------------------


def bind(object, method, function):
    """ Attaches the function as a method with the given name to the given object.
    """
    if new:
        # Python 2
        setattr(object, method, new.instancemethod(function, object))
    else:
        # Python 3: There is no good reason to use this function in Python 3.
        setattr(object, method, function)


class Stream(list):

    def __init__(self, url, delimiter="\n", **kwargs):
        """ Buffered stream of data from a given URL.
        """
        self.socket = URL(url).open(**kwargs)
        self.buffer = ""
        self.delimiter = delimiter

    def update(self, bytes=1024):
        """ Reads a number of bytes from the stream.
            If a delimiter is encountered, calls Stream.parse() on the packet.
        """
        packets = []
        self.buffer += self.socket.read(bytes).decode("utf-8")
        self.buffer = self.buffer.split(self.delimiter, 1)
        while len(self.buffer) > 1:
            data = self.buffer[0]
            data = self.parse(data)
            if data is not None:
                packets.append(data)
            self.buffer = self.buffer[-1]
            self.buffer = self.buffer.split(self.delimiter, 1)
        self.buffer = self.buffer[-1]
        self.extend(packets)
        return packets

    def parse(self, data):
        """ Must be overridden in a subclass.
        """
        return data

    def clear(self):
        list.__init__(self, [])


def stream(url, delimiter="\n", parse=lambda data: data, **kwargs):
    """ Returns a new Stream with the given parse method.
    """
    stream = Stream(url, delimiter, **kwargs)
    bind(stream, "parse", lambda stream, data: parse(data))
    return stream

#--- FIND URLs -------------------------------------------------------------------------------------
# Functions for parsing URL's and e-mail adresses from strings.

RE_URL_PUNCTUATION = ("\"'{(>", "\"'.,;)}")
RE_URL_HEAD = r"[%s|\[|\s]" % "|".join(RE_URL_PUNCTUATION[0])      # Preceded by space, parenthesis or HTML tag.
RE_URL_TAIL = r"[%s|\]]*[\s|\<]" % "|".join(RE_URL_PUNCTUATION[1]) # Followed by space, punctuation or HTML tag.
RE_URL1 = r"(https?://.*?)" + RE_URL_TAIL                          # Starts with http:// or https://
RE_URL2 = RE_URL_HEAD + r"(www\..*?\..*?)" + RE_URL_TAIL           # Starts with www.
RE_URL3 = RE_URL_HEAD + r"([\w|-]*?\.(com|net|org|edu|de|uk))" + RE_URL_TAIL

RE_URL1, RE_URL2, RE_URL3 = (
    re.compile(RE_URL1, re.I),
    re.compile(RE_URL2, re.I),
    re.compile(RE_URL3, re.I))


def find_urls(string, unique=True):
    """ Returns a list of URLs parsed from the string.
        Works on http://, https://, www. links or domain names ending in .com, .org, .net.
        Links can be preceded by leading punctuation (open parens)
        and followed by trailing punctuation (period, comma, close parens).
    """
    string = u(string)
    string = string.replace("\u2024", ".")
    string = string.replace(" ", "  ")
    matches = []
    for p in (RE_URL1, RE_URL2, RE_URL3):
        for m in p.finditer(" %s " % string):
            s = m.group(1)
            s = s.split("\">")[0].split("'>")[0] # google.com">Google => google.com
            if not unique or s not in matches:
                matches.append(s)
    return matches

links = find_urls

RE_EMAIL = re.compile(r"[\w\-\.\+]+@(\w[\w\-]+\.)+[\w\-]+") # tom.de+smedt@clips.ua.ac.be


def find_email(string, unique=True):
    """ Returns a list of e-mail addresses parsed from the string.
    """
    string = u(string).replace("\u2024", ".")
    matches = []
    for m in RE_EMAIL.finditer(string):
        s = m.group(0)
        if not unique or s not in matches:
            matches.append(s)
    return matches


def find_between(a, b, string):
    """ Returns a list of substrings between a and b in the given string.
    """
    p = "%s(.*?)%s" % (a, b)
    p = re.compile(p, re.DOTALL | re.I)
    return [m for m in p.findall(string)]

#### PLAIN TEXT ####################################################################################
# Functions for stripping HTML tags from strings.

BLOCK = [
    "title", "h1", "h2", "h3", "h4", "h5", "h6", "p",
    "center", "blockquote", "div", "table", "ul", "ol", "dl", "pre", "code", "form"
]

SELF_CLOSING = ["br", "hr", "img"]

# Element tag replacements for a stripped version of HTML source with strip_tags().
# Block-level elements are followed by linebreaks,
# list items are preceded by an asterisk ("*").
LIST_ITEM = "*"
blocks = dict.fromkeys(BLOCK + ["br", "tr", "td"], ("", "\n\n"))
blocks.update({
    "li": ("%s " % LIST_ITEM, "\n"),
   "img": ("", ""),
    "br": ("", "\n"),
    "th": ("", "\n"),
    "tr": ("", "\n"),
    "td": ("", "\t"),
})


class HTMLParser(_HTMLParser):

    def clean(self, html):
        html = decode_utf8(html)
        html = html.replace("/>", " />")
        html = html.replace("  />", " />")
        html = html.replace("<!", "&lt;!")
        html = html.replace("&lt;!DOCTYPE", "<!DOCTYPE")
        html = html.replace("&lt;!doctype", "<!doctype")
        html = html.replace("&lt;!--", "<!--")
        return html


class HTMLTagstripper(HTMLParser):

    def __init__(self):
        HTMLParser.__init__(self)

    def strip(self, html, exclude=[], replace=blocks):
        """ Returns the HTML string with all element tags (e.g. <p>) removed.
            - exclude    : a list of tags to keep. Element attributes are stripped.
                           To preserve attributes a dict of (tag name, [attribute])-items can be given.
            - replace    : a dictionary of (tag name, (replace_before, replace_after))-items.
                           By default, block-level elements are separated with linebreaks.
        """
        if html is None:
            return None
        self._exclude = isinstance(exclude, dict) and exclude or dict.fromkeys(exclude, [])
        self._replace = replace
        self._data = []
        self.feed(self.clean(html))
        self.close()
        self.reset()
        return "".join(self._data)

    def handle_starttag(self, tag, attributes):
        if tag in BLOCK and self._data and self._data[-1][-1:] != "\n":
            # Block-level elements always break to a new line.
            self._data.append("\n")
        if tag in self._exclude:
            # Create the tag attribute string,
            # including attributes defined in the HTMLTagStripper._exclude dict.
            a = len(self._exclude[tag]) > 0 and attributes or []
            a = ["%s=\"%s\"" % (k, v) for k, v in a if k in self._exclude[tag]]
            a = (" " + " ".join(a)).rstrip()
            self._data.append("<%s%s>" % (tag, a))
        if tag in self._replace:
            self._data.append(self._replace[tag][0])
        if tag in self._replace and tag in SELF_CLOSING:
            self._data.append(self._replace[tag][1])

    def handle_endtag(self, tag):
        if tag in self._exclude and self._data and self._data[-1].startswith("<" + tag):
            # Never keep empty elements (e.g. <a></a>).
            self._data.pop(-1)
            return
        if tag in self._exclude:
            self._data.append("</%s>" % tag)
        if tag in self._replace:
            self._data.append(self._replace[tag][1])

    def handle_data(self, data):
        self._data.append(data.strip("\n\t"))

    def handle_comment(self, comment):
        if "comment" in self._exclude or \
               "!--" in self._exclude:
            self._data.append("<!--%s-->" % comment)

# As a function:
strip_tags = HTMLTagstripper().strip


def strip_element(string, tag, attributes=""):
    """ Removes all elements with the given tagname and attributes from the string.
        Open and close tags are kept in balance.
        No HTML parser is used: strip_element(s, "a", 'class="x"') matches
        '<a class="x">' or '<a href="x" class="x">' but not "<a class='x'>".
    """
    s = string.lower() # Case-insensitive.
    t = tag.strip("</>")
    a = (" " + attributes.lower().strip()).rstrip()
    i = 0
    j = 0
    while j >= 0:
        #i = s.find("<%s%s" % (t, a), i)
        m = re.search(r"<%s[^\>]*?%s" % (t, a), s[i:])
        i = i + m.start() if m else -1
        j = s.find("</%s>" % t, i + 1)
        opened, closed = s[i:j].count("<%s" % t), 1
        while opened > closed and j >= 0:
            k = s.find("</%s>" % t, j + 1)
            opened += s[j:k].count("<%s" % t)
            closed += 1
            j = k
        if i < 0:
            return string
        if j < 0:
            return string[:i]
        string = string[:i] + string[j + len(t) + 3:]; s = string.lower()
    return string


def strip_between(a, b, string):
    """ Removes anything between (and including) string a and b inside the given string.
    """
    p = "%s.*?%s" % (a, b)
    p = re.compile(p, re.DOTALL | re.I)
    return re.sub(p, "", string)


def strip_javascript(html):
    return strip_between("<script.*?>", "</script>", html)


def strip_inline_css(html):
    return strip_between("<style.*?>", "</style>", html)


def strip_comments(html):
    return strip_between("<!--", "-->", html)


def strip_forms(html):
    return strip_between("<form.*?>", "</form>", html)

RE_AMPERSAND = re.compile("\&(?!\#)")           # & not followed by #
RE_UNICODE = re.compile(r'&(#?)(x|X?)(\w+);') # &#201;


def encode_entities(string):
    """ Encodes HTML entities in the given string ("<" => "&lt;").
        For example, to display "<em>hello</em>" in a browser,
        we need to pass "&lt;em&gt;hello&lt;/em&gt;" (otherwise "hello" in italic is displayed).
    """
    if isinstance(string, str):
        string = RE_AMPERSAND.sub("&amp;", string)
        string = string.replace("<", "&lt;")
        string = string.replace(">", "&gt;")
        string = string.replace('"', "&quot;")
        string = string.replace("'", "&#39;")
    return string


def decode_entities(string):
    """ Decodes HTML entities in the given string ("&lt;" => "<").
    """
    # http://snippets.dzone.com/posts/show/4569
    def replace_entity(match):
        hash, hex, name = match.group(1), match.group(2), match.group(3)
        if hash == "#" or name.isdigit():
            if hex == "":
                return chr(int(name))                 # "&#38;" => "&"
            if hex.lower() == "x":
                return chr(int("0x" + name, 16))      # "&#x0026;" = > "&"
        else:
            cp = name2codepoint.get(name) # "&amp;" => "&"
            return chr(cp) if cp else match.group()   # "&foo;" => "&foo;"
    if isinstance(string, str):
        return RE_UNICODE.subn(replace_entity, string)[0]
    return string


def encode_url(string):
    return quote_plus(bytestring(string)) # "black/white" => "black%2Fwhite".


def decode_url(string):
    return u(unquote_plus(bytestring(string)))

RE_SPACES = re.compile("( |\xa0)+", re.M) # Matches one or more spaces.
RE_TABS = re.compile(r"\t+", re.M)      # Matches one or more tabs.


def collapse_spaces(string, indentation=False, replace=" "):
    """ Returns a string with consecutive spaces collapsed to a single space.
        Whitespace on empty lines and at the end of each line is removed.
        With indentation=True, retains leading whitespace on each line.
    """
    p = []
    for x in string.splitlines():
        n = indentation and len(x) - len(x.lstrip()) or 0
        p.append(x[:n] + RE_SPACES.sub(replace, x[n:]).strip())
    return "\n".join(p)


def collapse_tabs(string, indentation=False, replace=" "):
    """ Returns a string with (consecutive) tabs replaced by a single space.
        Whitespace on empty lines and at the end of each line is removed.
        With indentation=True, retains leading whitespace on each line.
    """
    p = []
    for x in string.splitlines():
        n = indentation and len(x) - len(x.lstrip()) or 0
        p.append(x[:n] + RE_TABS.sub(replace, x[n:]).strip())
    return "\n".join(p)


def collapse_linebreaks(string, threshold=1):
    """ Returns a string with consecutive linebreaks collapsed to at most the given threshold.
        Whitespace on empty lines and at the end of each line is removed.
    """
    n = "\n" * threshold
    p = [s.rstrip() for s in string.splitlines()]
    string = "\n".join(p)
    string = re.sub(n + r"+", n, string)
    return string


def plaintext(html, keep=[], replace=blocks, linebreaks=2, indentation=False):
    """ Returns a string with all HTML tags removed.
        Content inside HTML comments, the <style> tag and the <script> tags is removed.
        - keep        : a list of tags to keep. Element attributes are stripped.
                        To preserve attributes a dict of (tag name, [attribute])-items can be given.
        - replace     : a dictionary of (tag name, (replace_before, replace_after))-items.
                        By default, block-level elements are followed by linebreaks.
        - linebreaks  : the maximum amount of consecutive linebreaks,
        - indentation : keep left line indentation (tabs and spaces)?
    """
    if isinstance(html, Element):
        html = html.content
    if not keep.__contains__("script"):
        html = strip_javascript(html)
    if not keep.__contains__("style"):
        html = strip_inline_css(html)
    if not keep.__contains__("form"):
        html = strip_forms(html)
    if not keep.__contains__("comment") and \
       not keep.__contains__("!--"):
        html = strip_comments(html)
    html = html.replace("\r", "\n")
    html = decode_entities(html)
    html = strip_tags(html, exclude=keep, replace=replace)
    html = collapse_spaces(html, indentation)
    html = collapse_tabs(html, indentation)
    html = collapse_linebreaks(html, linebreaks)
    html = html.strip()
    return html

#### SEARCH ENGINE #################################################################################

SEARCH    = "search"    # Query for pages (i.e. links to websites).
IMAGE     = "image"     # Query for images.
NEWS      = "news"      # Query for news items.

TINY      = "tiny"      # Image size around 100x100.
SMALL     = "small"     # Image size around 200x200.
MEDIUM    = "medium"    # Image size around 500x500.
LARGE     = "large"     # Image size around 1000x1000.

RELEVANCY = "relevancy" # Sort results by most relevant.
LATEST    = "latest"    # Sort results by most recent.


class Result(dict):

    def __init__(self, url, **kwargs):
        """ An item in a list of results returned by SearchEngine.search().
            All dictionary keys are available as Unicode string attributes.
            - id       : unique identifier,
            - url      : the URL of the referred web content,
            - title    : the title of the content at the URL,
            - text     : the content text,
            - language : the content language,
            - author   : for news items and posts, the author,
            - date     : for news items and posts, the publication date.
        """
        dict.__init__(self)
        self.url = url
        self.id = kwargs.pop("id", "")
        self.title = kwargs.pop("title", "")
        self.text = kwargs.pop("text", "")
        self.language = kwargs.pop("language", "")
        self.author = kwargs.pop("author", "")
        self.date = kwargs.pop("date", "")
        self.votes = kwargs.pop("votes", 0) # (e.g., Facebook likes)
        self.shares = kwargs.pop("shares", 0) # (e.g., Twitter retweets)
        self.comments = kwargs.pop("comments", 0)
        for k, v in kwargs.items():
            self[k] = v

    @property
    def txt(self):
        return self.text

    @property
    def description(self):
        return self.text # Backwards compatibility.

    @property
    def likes(self):
        return self.votes

    @property
    def retweets(self):
        return self.shares

    def download(self, *args, **kwargs):
        """ Download the content at the given URL.
            By default it will be cached - see URL.download().
        """
        return URL(self.url).download(*args, **kwargs)

    def _format(self, v):
        if isinstance(v, bytes): # Store strings as unicode.
            return u(v)
        if v is None:
            return ""
        return v

    def __getattr__(self, k):
        return self.get(k, "")

    def __getitem__(self, k):
        return self.get(k, "")

    def __setattr__(self, k, v):
        self.__setitem__(k, v)

    def __setitem__(self, k, v):
        dict.__setitem__(self, u(k), self._format(v))

    def setdefault(self, k, v=None):
        return dict.setdefault(self, u(k), self._format(v))

    def update(self, *args, **kwargs):
        dict.update(self, [(u(k), self._format(v)) for k, v in dict(*args, **kwargs).items()])

    def __repr__(self):
        return "Result(%s)" % repr(dict((k, v) for k, v in self.items() if v))


class Results(list):

    def __init__(self, source=None, query=None, type=SEARCH, total=0):
        """ A list of results returned from SearchEngine.search().
            - source: the service that yields the results (e.g. GOOGLE, TWITTER).
            - query : the query that yields the results.
            - type  : the query type (SEARCH, IMAGE, NEWS).
            - total : the total result count.
                      This is not the length of the list, but the total number of matches for the given query.
        """
        self.source = source
        self.query = query
        self.type = type
        self.total = total


class SearchEngine(object):

    def __init__(self, license=None, throttle=1.0, language=None):
        """ A base class for a web service.
            - license  : license key for the API,
            - throttle : delay between requests (avoid hammering the server).
            Inherited by: Google, Bing, Wikipedia, Twitter, Facebook, Flickr, ...
        """
        self.license = license
        self.throttle = throttle    # Amount of sleep time after executing a query.
        self.language = language    # Result.language restriction (e.g., "en").
        self.format = lambda x: x # Formatter applied to each attribute of each Result.

    def search(self, query, type=SEARCH, start=1, count=10, sort=RELEVANCY, size=None, cached=True, **kwargs):
        return Results(source=None, query=query, type=type)


class SearchEngineError(HTTPError):
    pass


class SearchEngineTypeError(SearchEngineError):
    pass # Raised when an unknown type is passed to SearchEngine.search().


class SearchEngineLimitError(SearchEngineError):
    pass # Raised when the query limit for a license is reached.

#--- GOOGLE ----------------------------------------------------------------------------------------
# Google Search is a web search engine owned by Google Inc.
# Google Custom Search is a paid service.
# https://code.google.com/apis/console/
# http://code.google.com/apis/customsearch/v1/overview.html

GOOGLE = "https://www.googleapis.com/customsearch/v1?"
GOOGLE_LICENSE = api.license["Google"]
GOOGLE_CUSTOM_SEARCH_ENGINE = "000579440470800426354:_4qo2s0ijsi"

# Search results can start with: "Jul 29, 2007 ...",
# which is the date of the page parsed by Google from the content.
RE_GOOGLE_DATE = re.compile("^([A-Z][a-z]{2} [0-9]{1,2}, [0-9]{4}) {0,1}...")


class Google(SearchEngine):

    def __init__(self, license=None, throttle=0.5, language=None):
        SearchEngine.__init__(self, license or GOOGLE_LICENSE, throttle, language)

    def search(self, query, type=SEARCH, start=1, count=10, sort=RELEVANCY, size=None, cached=True, **kwargs):
        """ Returns a list of results from Google for the given query.
            - type : SEARCH,
            - start: maximum 100 results => start 1-10 with count=10,
            - count: maximum 10,
            There is a daily limit of 10,000 queries. Google Custom Search is a paid service.
        """
        if type != SEARCH:
            raise SearchEngineTypeError
        if not query or count < 1 or start < 1 or start > (100 / count):
            return Results(GOOGLE, query, type)
        # 1) Create request URL.
        url = URL(GOOGLE, query={
              "key": self.license or GOOGLE_LICENSE,
               "cx": GOOGLE_CUSTOM_SEARCH_ENGINE,
                "q": query,
            "start": 1 + (start - 1) * count,
              "num": min(count, 10),
              "alt": "json"
        })
        # 2) Restrict language.
        if self.language is not None:
            url.query["lr"] = "lang_" + self.language
        # 3) Parse JSON response.
        kwargs.setdefault("unicode", True)
        kwargs.setdefault("throttle", self.throttle)
        data = url.download(cached=cached, **kwargs)
        data = json.loads(data)
        if data.get("error", {}).get("code") == 403:
            raise SearchEngineLimitError
        results = Results(GOOGLE, query, type)
        results.total = int(data.get("queries", {}).get("request", [{}])[0].get("totalResults") or 0)
        for x in data.get("items", []):
            r = Result(url=None)
            r.url      = self.format(x.get("link"))
            r.title    = self.format(x.get("title"))
            r.text     = self.format(x.get("htmlSnippet").replace("<br>  ", "").replace("<b>...</b>", "..."))
            r.language = self.language or ""
            r.date     = ""
            if not r.date:
                # Google Search results can start with a date (parsed from the content):
                m = RE_GOOGLE_DATE.match(r.text)
                if m:
                    r.date = m.group(1)
                    r.text = "..." + r.text[len(m.group(0)):]
            results.append(r)
        return results

    def translate(self, string, input="en", output="fr", **kwargs):
        """ Returns the translation of the given string in the desired output language.
            Google Translate is a paid service, license without billing raises HTTP401Authentication.
        """
        url = URL("https://www.googleapis.com/language/translate/v2?", method=GET, query={
               "key": self.license or GOOGLE_LICENSE,
                 "q": string.encode("utf-8"), # 1000 characters maximum
            "source": input,
            "target": output
        })
        kwargs.setdefault("cached", False)
        kwargs.setdefault("unicode", True)
        kwargs.setdefault("throttle", self.throttle)
        if input == output:
            return string
        try:
            data = url.download(**kwargs)
        except HTTP403Forbidden:
            raise HTTP401Authentication("Google translate API is a paid service")
        data = json.loads(data)
        data = data.get("data", {}).get("translations", [{}])[0].get("translatedText", "")
        data = decode_entities(data)
        return u(data)

    def identify(self, string, **kwargs):
        """ Returns a (language, confidence)-tuple for the given string.
            Google Translate is a paid service, license without billing raises HTTP401Authentication.
        """
        url = URL("https://www.googleapis.com/language/translate/v2/detect?", method=GET, query={
               "key": self.license or GOOGLE_LICENSE,
                 "q": string[:1000].encode("utf-8")
        })
        kwargs.setdefault("cached", False)
        kwargs.setdefault("unicode", True)
        kwargs.setdefault("throttle", self.throttle)
        try:
            data = url.download(**kwargs)
        except HTTP403Forbidden:
            raise HTTP401Authentication("Google translate API is a paid service")
        data = json.loads(data)
        data = data.get("data", {}).get("detections", [[{}]])[0][0]
        data = u(data.get("language")), float(data.get("confidence"))
        return data

#--- YAHOO -----------------------------------------------------------------------------------------
# Yahoo! Search is a web search engine owned by Yahoo! Inc.
# Yahoo! BOSS ("Build Your Own Search Service") is a paid service.
# http://developer.yahoo.com/search/

YAHOO = "http://yboss.yahooapis.com/ysearch/"
YAHOO_LICENSE = api.license["Yahoo"]


class Yahoo(SearchEngine):

    def __init__(self, license=None, throttle=0.5, language=None):
        SearchEngine.__init__(self, license or YAHOO_LICENSE, throttle, language)

    def _authenticate(self, url):
        url.query.update({
            "oauth_version": "1.0",
            "oauth_nonce": oauth.nonce(),
            "oauth_timestamp": oauth.timestamp(),
            "oauth_consumer_key": self.license[0],
            "oauth_signature_method": "HMAC-SHA1"
        })
        url.query["oauth_signature"] = oauth.sign(url.string.split("?")[0], url.query,
            method = url.method,
            secret = self.license[1]
        )
        return url

    def search(self, query, type=SEARCH, start=1, count=10, sort=RELEVANCY, size=None, cached=True, **kwargs):
        """ Returns a list of results from Yahoo for the given query.
            - type : SEARCH, IMAGE or NEWS,
            - start: maximum 1000 results => start 1-100 with count=10, 1000/count,
            - count: maximum 50, or 35 for images.
            There is no daily limit, however Yahoo BOSS is a paid service.
        """
        if type not in (SEARCH, IMAGE, NEWS):
            raise SearchEngineTypeError
        if type == SEARCH:
            url = YAHOO + "web"
        if type == IMAGE:
            url = YAHOO + "images"
        if type == NEWS:
            url = YAHOO + "news"
        if not query or count < 1 or start < 1 or start > 1000 / count:
            return Results(YAHOO, query, type)
        # 1) Create request URL.
        url = URL(url, method=GET, query={
                 "q": query.replace(" ", "+"),
             "start": 1 + (start - 1) * count,
             "count": min(count, type == IMAGE and 35 or 50),
            "format": "json"
        })
        # 2) Restrict language.
        if self.language is not None:
            market = locale.market(self.language)
            if market:
                url.query["market"] = market.lower()
        # 3) Authenticate.
        url = self._authenticate(url)
        # 4) Parse JSON response.
        kwargs.setdefault("unicode", True)
        kwargs.setdefault("throttle", self.throttle)
        try:
            data = url.download(cached=cached, **kwargs)
        except HTTP401Authentication:
            raise HTTP401Authentication("Yahoo %s API is a paid service" % type)
        except HTTP403Forbidden:
            raise SearchEngineLimitError
        data = json.loads(data)
        data = data.get("bossresponse") or {}
        data = data.get({SEARCH: "web", IMAGE: "images", NEWS: "news"}[type], {})
        results = Results(YAHOO, query, type)
        results.total = int(data.get("totalresults") or 0)
        for x in data.get("results", []):
            r = Result(url=None)
            r.url = self.format(x.get("url", x.get("clickurl")))
            r.title = self.format(x.get("title"))
            r.text = self.format(x.get("abstract"))
            r.date = self.format(x.get("date"))
            r.author = self.format(x.get("source"))
            r.language = self.format(x.get("language") and \
                                     x.get("language").split(" ")[0] or self.language or "")
            results.append(r)
        return results

#--- BING ------------------------------------------------------------------------------------------
# Bing is a web search engine owned by Microsoft.
# Bing Search API is a paid service.
# https://datamarket.azure.com/dataset/5BA839F1-12CE-4CCE-BF57-A49D98D29A44
# https://datamarket.azure.com/account/info

BING = "https://api.datamarket.azure.com/Bing/Search/"
BING_LICENSE = api.license["Bing"]


class Bing(SearchEngine):

    def __init__(self, license=None, throttle=0.5, language=None):
        SearchEngine.__init__(self, license or BING_LICENSE, throttle, language)

    def search(self, query, type=SEARCH, start=1, count=10, sort=RELEVANCY, size=None, cached=True, **kwargs):
        """" Returns a list of results from Bing for the given query.
             - type : SEARCH, IMAGE or NEWS,
             - start: maximum 1000 results => start 1-100 with count=10, 1000/count,
             - count: maximum 50, or 15 for news,
             - size : for images, either SMALL, MEDIUM or LARGE.
             There is no daily query limit.
        """
        if type not in (SEARCH, IMAGE, NEWS):
            raise SearchEngineTypeError
        if type == SEARCH:
            src = "Web"
        if type == IMAGE:
            src = "Image"
        if type == NEWS:
            src = "News"
        if not query or count < 1 or start < 1 or start > 1000 / count:
            return Results(BING + src + "?", query, type)
        # 1) Construct request URL.
        url = URL(BING + "Composite", method=GET, query={
               "Sources": "'" + src.lower() + "'",
                 "Query": "'" + query + "'",
                 "$skip": 1 + (start - 1) * count,
                  "$top": min(count, type == NEWS and 15 or 50),
               "$format": "json",
        })
        # 2) Restrict image size.
        if size in (TINY, SMALL, MEDIUM, LARGE):
            url.query["ImageFilters"] = {
                    TINY: "'Size:Small'",
                   SMALL: "'Size:Small'",
                  MEDIUM: "'Size:Medium'",
                   LARGE: "'Size:Large'"}[size]
        # 3) Restrict language.
        if type in (SEARCH, IMAGE) and self.language is not None:
            url.query["Query"] = url.query["Query"][:-1] + " language: %s'" % self.language
        if type in (NEWS,) and self.language is not None:
            market = locale.market(self.language)
            if market:
                url.query["Market"] = "'" + market + "'"
        # 4) Parse JSON response.
        kwargs["authentication"] = ("", self.license)
        kwargs.setdefault("unicode", True)
        kwargs.setdefault("throttle", self.throttle)
        try:
            data = url.download(cached=cached, **kwargs)
        except HTTP401Authentication:
            raise HTTP401Authentication("Bing %s API is a paid service" % type)
        except HTTP503ServiceUnavailable:
            raise SearchEngineLimitError
        data = json.loads(data)
        data = data.get("d", {})
        data = data.get("results", [{}])[0]
        results = Results(BING, query, type)
        results.total = int(data.get(src + "Total") or 0)
        for x in data.get(src, []):
            r = Result(url=None)
            r.url = self.format(x.get("MediaUrl", x.get("Url")))
            r.title = self.format(x.get("Title"))
            r.text = self.format(x.get("Description", x.get("Snippet")))
            r.language = self.language or ""
            r.date = self.format(x.get("DateTime", x.get("Date")))
            r.author = self.format(x.get("Source"))
            results.append(r)
        return results

#--- DUCKDUCKGO ------------------------------------------------------------------------------------
# DuckDuckGo is a privacy-respecting aggregate search engine,
# with information from Wikipedia, WikiHow, Wikia, GitHub, The Free Dictionary, etc.
# https://duckduckgo.com/api.html
# https://duckduckgo.com/params.html

DUCKDUCKGO = "http://api.duckduckgo.com/"
DUCKDUCKGO_LICENSE = api.license["DuckDuckGo"]

# Results from DuckDuckGo have a Result.type with semantic information,
# e.g., "apple" => "plant and plant parts". Known types:
REFERENCE, CATEGORY, DEFINITION = \
    "reference", "category", "definition"


class DuckDuckGo(SearchEngine):

    def __init__(self, license=None, throttle=0.5, language=None):
        SearchEngine.__init__(self, license or DUCKDUCKGO_LICENSE, throttle, language)

    def search(self, query, type=SEARCH, start=None, count=None, sort=RELEVANCY, size=None, cached=True, **kwargs):
        """" Returns a list of results from DuckDuckGo for the given query.
        """
        if type != SEARCH:
            raise SearchEngineTypeError
        # 1) Construct request URL.
        url = URL(DUCKDUCKGO, method=GET, query={
            "q": query,
            "o": "json"
        })
        # 2) Restrict language.
        if type == SEARCH and self.language is not None:
            url.query["kl"] = self.language
        # 3) Parse JSON response.
        kwargs.setdefault("unicode", True)
        kwargs.setdefault("throttle", self.throttle)
        data = url.download(cached=cached, **kwargs)
        data = json.loads(data)
        results = Results(DUCKDUCKGO, query, type)
        results.total = None
        for x in data.get("Results", []):
            if x.get("FirstURL"):
                r = Result(url=None)
                # Parse official website link.
                r.url = self.format(x.get("FirstURL"))
                r.title = self.format(data.get("Heading"))
                r.text = self.format(data.get("Abstract"))
                r.author = self.format(data.get("AbstractSource"))
                r.type = self.format(REFERENCE)
                results.append(r)
        for topic in data.get("RelatedTopics", []):
            for x in topic.get("Topics", [topic]):
                r = Result(url=None)
                r.url = x.get("FirstURL")
                # Parse title and type from URL (e.g., http://duckduckgo.com/d/Cats?kl=en).
                m = re.match(r"^http://duckduckgo.com/([a-z]/)?(.*?)(\?|$)", r.url)
                # Parse title: "Cats".
                s1 = m and m.group(2) or "" # Title: "Cats"
                s1 = u(decode_url(s1.encode("utf-8")))
                s1 = s1.strip().replace("_", " ")
                s1 = s1[:1].upper() + s1[1:]
                # Parse description; the part before the first "-" or "," was the link.
                s2 = x.get("Text", "").strip()
                s2 = re.sub(r" +", " ", s2)
                s2 = s2[:1].upper() + s2[1:] or ""
                s2 = s2.startswith(s1) \
                    and "<a href=\"%s\">%s</a>%s" % (r.url, s1, s2[len(s1):]) \
                     or re.sub(r"^(.*?)( - | or |, )(.*?)", "<a href=\"%s\">\\1</a>\\2\\3" % r.url, s2)
                # Parse type: "d/" => "definition".
                s3 = m and m.group(1) or ""
                s3 = {"c": CATEGORY, "d": DEFINITION}.get(s3.rstrip("/"), "")
                s3 = topic.get("Name", "").lower() or s3
                s3 = re.sub("^in ", "", s3)
                # Format result.
                r.url = self.format(r.url)
                r.title = self.format(s1)
                r.text = self.format(s2)
                r.type = self.format(s3)
                results.append(r)
        return results

    def answer(self, string, **kwargs):
        """ Returns a DuckDuckGo answer for the given string (e.g., math, spelling, ...)
        """
        url = URL(DUCKDUCKGO, method=GET, query={
            "q": string.encode("utf-8"),
            "o": "json"
        })
        kwargs.setdefault("cached", False)
        kwargs.setdefault("unicode", True)
        kwargs.setdefault("throttle", self.throttle)
        data = url.download(**kwargs)
        data = json.loads(data)
        data = data.get(kwargs.get("field", "Answer"))
        return u(data)

    def spelling(self, string):
        """ Returns a list of spelling suggestions for the given string.
        """
        s = self.answer("spell " + string, cached=True)
        s = re.findall(r"<a.*?>(.*?)</a>", s)
        return s

    def definition(self, string):
        """ Returns a dictionary definition for the given string.
        """
        s = self.answer(string, field="Definition", cached=True)
        s = re.sub(r"^.*? definition: ", "", s)
        s = re.sub(r"(^'''.*?''' |^)(.)(.*?)$",
            lambda m: m.group(1) + m.group(2).upper() + m.group(3), s)
        return s

DDG = DuckDuckGo

#for r in DDG().search("cats"):
#    print(r.url)
#    print(r.title) # Can be used as a new query.
#    print(plaintext(r.text))
#    print(r.type)  # REFERENCE, CATEGORY, DEFINITION, "people", "sports" ...
#    print()

#print(DDG().definition("cat"))
#print(DDG().spelling("catnpa"))

#--- FAROO
#-----------------------------------------------------------------------------------------
# Faroo is web search engine with 1 million free queries per month.
# http://www.faroo.com/hp/api/api.html

FAROO = "http://www.faroo.com/api"
FAROO_LICENSE = api.license["Faroo"]


class Faroo(SearchEngine):

    def __init__(self, license = None, throttle = 0.5, language = None):
        SearchEngine.__init__(self, license or FAROO_LICENSE, throttle, language)

    def search(self, query, type = SEARCH, start = 1, count = 10, sort = RELEVANCY, size = None, cached = True, **kwargs):
        if type != SEARCH:
            raise SearchEngineTypeError

        # 1) Create request URL.
        url = URL(FAROO, method=GET, query={
                 "q" : query.replace(" ", "+"),
             "start" : 1 + (start - 1) * count,
            "length" : count,
               "key" : self.license or FAROO_LICENSE,
                 "f" : "json"
        })
        # 2) Restrict language
        if self.language in ('en', 'de', 'zh'):
            url.query["language"] = self.language
        else:
            raise SearchEngineError("Language not supported")
        # 3) Parse JSON response.
        kwargs.setdefault("unicode", True)
        kwargs.setdefault("throttle", self.throttle)
        try:
            data = url.download(cached = cached, user_agent = "*", **kwargs)
        except HTTP401Authentication:
            raise HTTP401Authentication("Faroo %s API requires an API key" % type)
        except HTTP403Forbidden:
            raise SearchEngineLimitError
        data = json.loads(data)
        results = Results(FAROO, query, type)
        results.total = int(data.get("count") or 0)
        for x in data.get("results", []):
            r = Result(url=None)
            r.url = self.format(x.get("url"))
            r.title = self.format(x.get("title"))
            r.text = self.format(x.get("kwic"))
            r.date = self.format(x.get("date"))
            r.author = self.format(x.get("author"))
            r.language = self.format(self.language or "")
            results.append(r)
        return results

#--- TWITTER ---------------------------------------------------------------------------------------
# Twitter is an online social networking service and microblogging service,
# that enables users to post and read text-based messages of up to 140 characters ("tweets").
# https://dev.twitter.com/docs/api/1.1

TWITTER         = "https://api.twitter.com/1.1/"
TWITTER_STREAM  = "https://stream.twitter.com/1.1/statuses/filter.json"
TWITTER_STATUS  = "https://twitter.com/%s/status/%s"
TWITTER_LICENSE = api.license["Twitter"]
TWITTER_HASHTAG = re.compile(r"(\s|^)(#[a-z0-9_\-]+)", re.I)    # Word starts with "#".
TWITTER_MENTION = re.compile(r"(\s|^)(@[a-z0-9_\-]+)", re.I)    # Word starts with "@".
TWITTER_RETWEET = re.compile(r"(\s|^RT )(@[a-z0-9_\-]+)", re.I) # Word starts with "RT @".


class Twitter(SearchEngine):

    def __init__(self, license=None, throttle=0.5, language=None):
        SearchEngine.__init__(self, license or TWITTER_LICENSE, throttle, language)
        self._pagination = {}

    def _authenticate(self, url):
        url.query.update({
            "oauth_version": "1.0",
            "oauth_nonce": oauth.nonce(),
            "oauth_timestamp": oauth.timestamp(),
            "oauth_consumer_key": self.license[0],
            "oauth_token": self.license[2][0],
            "oauth_signature_method": "HMAC-SHA1"
        })
        url.query["oauth_signature"] = oauth.sign(url.string.split("?")[0], url.query,
            method = url.method,
            secret = self.license[1],
             token = self.license[2][1]
        )
        return url

    def search(self, query, type=SEARCH, start=1, count=10, sort=RELEVANCY, size=None, cached=False, **kwargs):
        """ Returns a list of results from Twitter for the given query.
            - type : SEARCH,
            - start: Result.id or int,
            - count: maximum 100.
            There is a limit of 150+ queries per 15 minutes.
        """
        if type != SEARCH:
            raise SearchEngineTypeError
        if not query or count < 1 or (isinstance(start, (int, float)) and start < 1):
            return Results(TWITTER, query, type)
        if not isinstance(start, (int, float)):
            id = int(start) - 1 if start and start.isdigit() else ""
        else:
            if start == 1:
                self._pagination = {}
            if start <= 10000:
                id = (query, kwargs.get("geo"), kwargs.get("date"), int(start) - 1, count)
                id = self._pagination.get(id, "")
            else:
                id = int(start) - 1
        # 1) Construct request URL.
        url = URL(TWITTER + "search/tweets.json?", method=GET)
        url.query = {
               "q": query,
          "max_id": id,
           "count": min(count, 100)
        }
        # 2) Restrict location with geo=(latitude, longitude, radius).
        #    It can also be a (latitude, longitude)-tuple with default radius "10km".
        if "geo" in kwargs:
            url.query["geocode"] = ",".join((map(str, kwargs.pop("geo")) + ["10km"])[:3])
        # 3) Restrict most recent with date="YYYY-MM-DD".
        #    Only older tweets are returned.
        if "date" in kwargs:
            url.query["until"] = kwargs.pop("date")
        # 4) Restrict language.
        url.query["lang"] = self.language or ""
        # 5) Authenticate.
        url = self._authenticate(url)
        # 6) Parse JSON response.
        kwargs.setdefault("unicode", True)
        kwargs.setdefault("throttle", self.throttle)
        try:
            data = url.download(cached=cached, **kwargs)
        except HTTP420Error:
            raise SearchEngineLimitError
        except HTTP429TooMayRequests:
            raise SearchEngineLimitError
        data = json.loads(data)
        results = Results(TWITTER, query, type)
        results.total = None
        for x in data.get("statuses", []):
            r = Result(url=None)
            r.id = self.format(x.get("id_str"))
            r.url = self.format(TWITTER_STATUS % (x.get("user", {}).get("screen_name"), x.get("id_str")))
            r.text = self.format(x.get("text"))
            r.date = self.format(x.get("created_at"))
            r.author = self.format(x.get("user", {}).get("screen_name"))
            r.language = self.format(x.get("metadata", {}).get("iso_language_code"))
            r.shares = self.format(x.get("retweet_count", 0))
            r.profile = self.format(x.get("user", {}).get("profile_image_url")) # Profile picture URL.
            # Fetch original status if retweet is truncated (i.e., ends with "...").
            rt = x.get("retweeted_status", None)
            if rt:
                comment = re.search(r"^(.*? )RT", r.text)
                comment = comment.group(1) if comment else ""
                r.text = self.format("RT @%s: %s" % (rt["user"]["screen_name"], rt["text"]))
            results.append(r)
        # Twitter.search(start=id, count=10) takes a tweet.id,
        # and returns 10 results that are older than this id.
        # In the past, start took an int used for classic pagination.
        # However, new tweets may arrive quickly,
        # so that by the time Twitter.search(start=2) is called,
        # it will yield results from page 1 (or even newer results).
        # For backward compatibility, we keep page cache,
        # that remembers the last id for a "page" for a given query,
        # when called in a loop.
        #
        # Store the last id retrieved.
        # If search() is called again with start+1, start from this id.
        if isinstance(start, (int, float)):
            k = (query, kwargs.get("geo"), kwargs.get("date"), int(start), count)
            if results:
                self._pagination[k] = str(int(results[-1].id) - 1)
            else:
                self._pagination[k] = id
        return results

    def profile(self, query, start=1, count=10, **kwargs):
        """ Returns a list of results for the given author id, alias or search query.
        """
        # 1) Construct request URL.
        url = URL(TWITTER + "users/search.json?", method=GET, query={
               "q": query,
            "page": start,
           "count": count
        })
        url = self._authenticate(url)
        # 2) Parse JSON response.
        kwargs.setdefault("cached", True)
        kwargs.setdefault("unicode", True)
        kwargs.setdefault("throttle", self.throttle)
        try:
            data = URL(url).download(**kwargs)
            data = json.loads(data)
        except HTTP400BadRequest:
            return []
        return [
            Result(url = "https://www.twitter.com/" + x.get("screen_name", ""),
                    id = x.get("id_str", ""),              # 14898655
                handle = x.get("screen_name", ""),         # tom_de_smedt
                  name = x.get("name", ""),                # Tom De Smedt
                  text = x.get("description", ""),         # Artist, scientist, software engineer
              language = x.get("lang", ""),                # en
                  date = x.get("created_at"),              # Sun May 10 10:00:00
                locale = x.get("location", ""),            # Belgium
               picture = x.get("profile_image_url", ""),   # http://pbs.twimg.com/...
               friends = int(x.get("followers_count", 0)), # 100
                 posts = int(x.get("statuses_count", 0))   # 100
            ) for x in data
        ]

    def trends(self, **kwargs):
        """ Returns a list with 10 trending topics on Twitter.
        """
        # 1) Construct request URL.
        url = URL("https://api.twitter.com/1.1/trends/place.json?id=1")
        url = self._authenticate(url)
        # 2) Parse JSON response.
        kwargs.setdefault("cached", False)
        kwargs.setdefault("unicode", True)
        kwargs.setdefault("throttle", self.throttle)
        try:
            data = url.download(**kwargs)
            data = json.loads(data)
        except HTTP400BadRequest:
            return []
        return [u(x.get("name")) for x in data[0].get("trends", [])]

    def stream(self, query, **kwargs):
        """ Returns a live stream of Result objects for the given query.
        """
        url = URL(TWITTER_STREAM)
        url.query["track"] = query
        url = self._authenticate(url)
        return TwitterStream(url, delimiter="\n", format=self.format, **kwargs)


class TwitterStream(Stream):

    def __init__(self, socket, delimiter="\n", format=lambda s: s, **kwargs):
        kwargs.setdefault("timeout", 30)
        Stream.__init__(self, socket, delimiter, **kwargs)
        self.format = format

    def parse(self, data):
        """ TwitterStream.queue will populate with Result objects as
            TwitterStream.update() is called iteratively.
        """
        if data.strip():
            x = json.loads(data)
            r = Result(url=None)
            r.id = self.format(x.get("id_str"))
            r.url = self.format(TWITTER_STATUS % (x.get("user", {}).get("screen_name"), x.get("id_str")))
            r.text = self.format(x.get("text"))
            r.date = self.format(x.get("created_at"))
            r.author = self.format(x.get("user", {}).get("screen_name"))
            r.language = self.format(x.get("metadata", {}).get("iso_language_code"))
            r.shares = self.format(x.get("retweet_count", 0))
            r.profile = self.format(x.get("user", {}).get("profile_image_url")) # Profile picture URL.
            # Fetch original status if retweet is truncated (i.e., ends with "...").
            rt = x.get("retweeted_status", None)
            if rt:
                comment = re.search(r"^(.*? )RT", r.text)
                comment = comment.group(1) if comment else ""
                r.text = self.format("RT @%s: %s" % (rt["user"]["screen_name"], rt["text"]))
            return r


def author(name):
    """ Returns a Twitter query-by-author-name that can be passed to Twitter.search().
        For example: Twitter().search(author("tom_de_smedt"))
    """
    return "from:%s" % name


def hashtags(string):
    """ Returns a list of hashtags (words starting with a #hash) from a tweet.
    """
    return [b for a, b in TWITTER_HASHTAG.findall(string)]


def mentions(string):
    """ Returns a list of mentions (words starting with a @author) from a tweet.
    """
    return [b for a, b in TWITTER_MENTION.findall(string)]


def retweets(string):
    """ Returns a list of retweets (words starting with a RT @author) from a tweet.
    """
    return [b for a, b in TWITTER_RETWEET.findall(string)]

#engine = Twitter()
#for i in range(2):
#    for tweet in engine.search("cat nap", cached=False, start=i+1, count=10):
#        print()
#        print(tweet.id)
#        print(tweet.url)
#        print(tweet.text)
#        print(tweet.author)
#        print(tweet.profile)
#        print(tweet.language)
#        print(tweet.date)
#        print(hashtags(tweet.text))
#        print(retweets(tweet.text))

#stream = Twitter().stream("cat")
#for i in range(10):
#    print(i)
#    stream.update()
#    for tweet in reversed(stream):
#        print(tweet.id)
#        print(tweet.text)
#        print(tweet.url)
#        print(tweet.language)
#    print()
#stream.clear()

#--- MEDIAWIKI -------------------------------------------------------------------------------------
# MediaWiki is a free wiki software application.
# MediaWiki powers popular websites such as Wikipedia, Wiktionary and Wikia.
# http://www.mediawiki.org/wiki/API:Main_page
# http://en.wikipedia.org/w/api.php

WIKIA = "http://wikia.com"
WIKIPEDIA = "http://wikipedia.com"
WIKIPEDIA_LICENSE = api.license["Wikipedia"]
MEDIAWIKI_LICENSE = None
MEDIAWIKI = "http://{SUBDOMAIN}.{DOMAIN}{API}"

# Pattern for meta links (e.g. Special:RecentChanges).
# http://en.wikipedia.org/wiki/Main_namespace
MEDIAWIKI_NAMESPACE = ["Main", "User", "Wikipedia", "File", "MediaWiki", "Template", "Help", "Category", "Portal", "Book"]
MEDIAWIKI_NAMESPACE += [s + " talk" for s in MEDIAWIKI_NAMESPACE] + ["Talk", "Special", "Media"]
MEDIAWIKI_NAMESPACE += ["WP", "WT", "MOS", "C", "CAT", "Cat", "P", "T", "H", "MP", "MoS", "Mos"]
_mediawiki_namespace = re.compile(r"^(" + "|".join(MEDIAWIKI_NAMESPACE) + "):", re.I)

# Pattern to identify disambiguation pages.
MEDIAWIKI_DISAMBIGUATION = "<a href=\"/wiki/Help:Disambiguation\" title=\"Help:Disambiguation\">disambiguation</a> page"

# Pattern to identify references, e.g. [12]
MEDIAWIKI_REFERENCE = r"\s*\[[0-9]{1,3}\]"

# Mediawiki.search(type=ALL).
ALL = "all"


class MediaWiki(SearchEngine):

    def __init__(self, license=None, throttle=5.0, language="en"):
        SearchEngine.__init__(self, license or MEDIAWIKI_LICENSE, throttle, language)

    @property
    def _url(self):
        # Must be overridden in a subclass; see Wikia and Wikipedia.
        return None

    @property
    def MediaWikiArticle(self):
        return MediaWikiArticle

    @property
    def MediaWikiSection(self):
        return MediaWikiSection

    @property
    def MediaWikiTable(self):
        return MediaWikiTable

    def __iter__(self):
        return self.articles()

    def articles(self, **kwargs):
        """ Returns an iterator over all MediaWikiArticle objects.
            Optional parameters can include those passed to
            MediaWiki.index(), MediaWiki.search() and URL.download().
        """
        for title in self.index(**kwargs):
            yield self.search(title, **kwargs)

    # Backwards compatibility.
    all = articles

    def index(self, namespace=0, start=None, count=100, cached=True, **kwargs):
        """ Returns an iterator over all article titles (for a given namespace id).
        """
        kwargs.setdefault("unicode", True)
        kwargs.setdefault("throttle", self.throttle)
        # Fetch article titles (default) or a custom id.
        id = kwargs.pop("_id", "title")
        # Loop endlessly (= until the last request no longer yields an "apcontinue").
        # See: http://www.mediawiki.org/wiki/API:Allpages
        while start != -1:
            url = URL(self._url, method=GET, query={
                     "action": "query",
                       "list": "allpages",
                "apnamespace": namespace,
                     "apfrom": start or "",
                    "aplimit": min(count, 500),
              "apfilterredir": "nonredirects",
                     "format": "json"
            })
            data = url.download(cached=cached, **kwargs)
            data = json.loads(data)
            for x in data.get("query", {}).get("allpages", {}):
                if x.get(id):
                    yield x[id]
            start = data.get("query-continue", {}).get("allpages", {})
            start = start.get("apcontinue", start.get("apfrom", -1))
        raise StopIteration

    # Backwards compatibility.
    list = index

    def search(self, query, type=SEARCH, start=1, count=10, sort=RELEVANCY, size=None, cached=True, **kwargs):
        """ With type=SEARCH, returns a MediaWikiArticle for the given query (case-sensitive).
            With type=ALL, returns a list of results. 
            Each result.title is the title of an article that contains the given query.
        """
        if type not in (SEARCH, ALL, "*"):
            raise SearchEngineTypeError
        if type == SEARCH: # Backwards compatibility.
            return self.article(query, cached=cached, **kwargs)
        if not query or start < 1 or count < 1:
            return Results(self._url, query, type)
        # 1) Construct request URL (e.g., Wikipedia for a given language).
        url = URL(self._url, method=GET, query={
            "action": "query",
              "list": "search",
          "srsearch": query,
          "sroffset": (start - 1) * count,
           "srlimit": min(count, 100),
            "srprop": "snippet",
            "format": "json"
        })
        # 2) Parse JSON response.
        kwargs.setdefault("unicode", True)
        kwargs.setdefault("throttle", self.throttle)
        data = url.download(cached=cached, **kwargs)
        data = json.loads(data)
        data = data.get("query", {})
        results = Results(self._url, query, type)
        results.total = int(data.get("searchinfo", {}).get("totalhits", 0))
        for x in data.get("search", []):
            u = "http://%s/wiki/%s" % (URL(self._url).domain, x.get("title").replace(" ", "_"))
            r = Result(url=u)
            r.id = self.format(x.get("title"))
            r.title = self.format(x.get("title"))
            r.text = self.format(plaintext(x.get("snippet")))
            results.append(r)
        return results

    def article(self, query, cached=True, **kwargs):
        """ Returns a MediaWikiArticle for the given query.
            The query is case-sensitive, for example on Wikipedia:
            - "tiger" = Panthera tigris,
            - "TIGER" = Topologically Integrated Geographic Encoding and Referencing.
        """
        url = URL(self._url, method=GET, query={
            "action": "parse",
              "page": query.replace(" ", "_"),
         "redirects": 1,
            "format": "json"
        })
        kwargs.setdefault("unicode", True)
        kwargs.setdefault("timeout", 30) # Parsing the article takes some time.
        kwargs.setdefault("throttle", self.throttle)
        data = url.download(cached=cached, **kwargs)
        data = json.loads(data)
        data = data.get("parse", {})
        a = self._parse_article(data, query=query)
        a = self._parse_article_sections(a, data)
        a = self._parse_article_section_structure(a)
        if not a.html or "id=\"noarticletext\"" in a.html:
            return None
        return a

    def _parse_article(self, data, **kwargs):
        return self.MediaWikiArticle(
                  title = plaintext(data.get("displaytitle", data.get("title", ""))),
                 source = data.get("text", {}).get("*", ""),
         disambiguation = data.get("text", {}).get("*", "").find(MEDIAWIKI_DISAMBIGUATION) >= 0,
                  links = [x["*"] for x in data.get("links", []) if not _mediawiki_namespace.match(x["*"])],
             categories = [x["*"] for x in data.get("categories", [])],
               external = [x for x in data.get("externallinks", [])],
                  media = [x for x in data.get("images", [])],
              redirects = [x for x in data.get("redirects", [])],
              languages = dict([(x["lang"], x["*"]) for x in data.get("langlinks", [])]),
              language = self.language,
                 parser = self, **kwargs)

    def _parse_article_sections(self, article, data):
        # If "References" is a section in the article,
        # the HTML will contain a marker <h*><span class="mw-headline" id="References">.
        # http://en.wikipedia.org/wiki/Section_editing
        t = article.title
        d = 0
        i = 0
        for x in data.get("sections", {}):
            a = x.get("anchor")
            if a:
                p = r"<h.>\s*.*?\s*<span class=\"mw-headline\" id=\"%s\">" % a
                p = re.compile(p)
                m = p.search(article.source, i)
                if m:
                    j = m.start()
                    t = plaintext(t)
                    article.sections.append(self.MediaWikiSection(article,
                        title = t,
                        start = i,
                         stop = j,
                        level = d))
                    t = plaintext(x.get("line", ""))
                    d = int(x.get("level", 2)) - 1
                    i = j
        return article

    def _parse_article_section_structure(self, article):
        # Sections with higher level are children of previous sections with lower level.
        for i, s2 in enumerate(article.sections):
            for s1 in reversed(article.sections[:i]):
                if s1.level < s2.level:
                    s2.parent = s1
                    s1.children.append(s2)
                    break
        return article


class MediaWikiArticle(object):

    def __init__(self, title="", source="", links=[], categories=[], languages={}, disambiguation=False, **kwargs):
        """ A MediaWiki article returned from MediaWiki.search().
            MediaWikiArticle.string contains the HTML content.
        """
        self.title          = title          # Article title.
        self.source         = source         # Article HTML content.
        self.sections       = []             # Article sections.
        self.links          = links          # List of titles of linked articles.
        self.categories     = categories     # List of categories. As links, prepend "Category:".
        self.external       = []             # List of external links.
        self.media          = []             # List of linked media (images, sounds, ...)
        self.disambiguation = disambiguation # True when the article is a disambiguation page.
        self.languages      = languages      # Dictionary of (language, article)-items, e.g. Cat => ("nl", "Kat")
        self.language       = kwargs.get("language", "en")
        self.redirects      = kwargs.get("redirects", [])
        self.parser         = kwargs.get("parser", MediaWiki())
        for k, v in kwargs.items():
            setattr(self, k, v)

    def _plaintext(self, string, **kwargs):
        """ Strips HTML tags, whitespace and wiki markup from the HTML source, including:
            metadata, info box, table of contents, annotations, thumbnails, disambiguation link.
            This is called internally from MediaWikiArticle.string.
        """
        s = string
        # Strip meta <table> elements.
        s = strip_element(s, "table", "id=\"toc")             # Table of contents.
        s = strip_element(s, "table", "class=\"infobox")      # Infobox.
        s = strip_element(s, "table", "class=\"navbox")       # Navbox.
        s = strip_element(s, "table", "class=\"mbox")         # Message.
        s = strip_element(s, "table", "class=\"metadata")     # Metadata.
        s = strip_element(s, "table", "class=\".*?wikitable") # Table.
        s = strip_element(s, "table", "class=\"toc")          # Table (usually footer).
        # Strip meta <div> elements.
        s = strip_element(s, "div", "id=\"toc")               # Table of contents.
        s = strip_element(s, "div", "class=\"infobox")        # Infobox.
        s = strip_element(s, "div", "class=\"navbox")         # Navbox.
        s = strip_element(s, "div", "class=\"mbox")           # Message.
        s = strip_element(s, "div", "class=\"metadata")       # Metadata.
        s = strip_element(s, "div", "id=\"annotation")        # Annotations.
        s = strip_element(s, "div", "class=\"dablink")        # Disambiguation message.
        s = strip_element(s, "div", "class=\"magnify")        # Thumbnails.
        s = strip_element(s, "div", "class=\"thumb ")         # Thumbnail captions.
        s = strip_element(s, "div", "class=\"barbox")         # Bar charts.
        s = strip_element(s, "div", "class=\"noprint")        # Hidden from print.
        s = strip_element(s, "sup", "class=\"noprint")
        # Strip absolute elements (don't know their position).
        s = strip_element(s, "div", "style=\"position:absolute")
        # Strip meta <span> elements.
        s = strip_element(s, "span", "class=\"error")
        # Strip math formulas, add [math] placeholder.
        s = re.sub(r"<img class=\"tex\".*?/>", "[math]", s)   # LaTex math images.
        s = plaintext(s, **kwargs)
        # Strip [edit] link (language dependent.)
        s = re.sub(r"\[edit\]\s*", "", s)
        s = re.sub(r"\[%s\]\s*" % {
            "en": "edit",
            "es": "editar código",
            "de": "Bearbeiten",
            "fr": "modifier le code",
            "it": "modifica sorgente",
            "nl": "bewerken",
        }.get(self.language, "edit"), "", s)
        # Insert space before inline references.
        s = s.replace("[", " [").replace("  [", " [")
        # Strip inline references.
        #s = re.sub(r" \[[0-9]+\]", "", s)
        return s

    def plaintext(self, **kwargs):
        return self._plaintext(self.source, **kwargs)

    @property
    def html(self):
        return self.source

    @property
    def src(self):
        return self.source

    @property
    def string(self):
        return self.plaintext()

    def __repr__(self):
        return "MediaWikiArticle(title=%s)" % repr(self.title)


class MediaWikiSection(object):

    def __init__(self, article, title="", start=0, stop=0, level=1):
        """ A (nested) section in the content of a MediaWikiArticle.
        """
        self.article  = article # MediaWikiArticle the section is part of.
        self.parent   = None    # MediaWikiSection the section is part of.
        self.children = []      # MediaWikiSections belonging to this section.
        self.title    = title   # Section title.
        self._start   = start   # Section start index in MediaWikiArticle.string.
        self._stop    = stop    # Section stop index in MediaWikiArticle.string.
        self._level   = level   # Section depth (main title + intro = level 0).
        self._links   = None
        self._tables  = None

    def plaintext(self, **kwargs):
        return self.article._plaintext(self.source, **kwargs)

    @property
    def source(self):
        return self.article.source[self._start:self._stop]

    @property
    def html(self):
        return self.source

    @property
    def src(self):
        return self.source

    @property
    def string(self):
        return self.plaintext()

    @property
    def content(self):
        # ArticleSection.string, minus the title.
        s = self.plaintext()
        t = plaintext(self.title)
        if s == t or (len(s) > len(t)) and s.startswith(t) and s[len(t)] not in string.punctuation + " ":
            return s[len(t):].lstrip()
        return s

    @property
    def links(self, path="/wiki/"):
        """ Yields a list of Wikipedia links in this section. Similar
            in functionality to MediaWikiArticle.links.
        """
        if self._links is None:
            a = HTMLLinkParser().parse(self.source)
            a = (decode_url(a.url) for a in a)
            a = (a[len(path):].replace("_", " ") for a in a if a.startswith(path))
            a = (a for a in a if not _mediawiki_namespace.match(a))
            self._links = sorted(set(a))
        return self._links

    @property
    def tables(self):
        """ Yields a list of MediaWikiTable objects in the section.
        """
        if self._tables is None:
            self._tables = []
            for style in ("wikitable", "sortable wikitable"):
                b = "<table class=\"%s\"" % style, "</table>"
                p = self.article._plaintext
                f = find_between
                for s in f(b[0], b[1], self.source):
                    t = self.article.parser.MediaWikiTable(self,
                         title = p((f(r"<caption.*?>", "</caption>", s) + [""])[0]),
                        source = b[0] + s + b[1])
                    # 1) Parse <td> and <th> content and format it as plain text.
                    # 2) Parse <td colspan=""> attribute, duplicate spanning cells.
                    # 3) For <th> in the first row, update MediaWikiTable.headers.
                    for i, row in enumerate(f(r"<tr", "</tr>", s)):
                        r1 = f(r"<t[d|h]", r"</t[d|h]>", row)
                        r1 = (((f(r'colspan="', r'"', v) + [1])[0], v[v.find(">") + 1:]) for v in r1)
                        r1 = ((int(n), v) for n, v in r1)
                        r2 = []
                        [[r2.append(p(v)) for j in range(n)] for n, v in r1]
                        if i == 0 and "</th>" in row:
                            t.headers = r2
                        else:
                            t.rows.append(r2)
                    self._tables.append(t)
        return self._tables

    @property
    def level(self):
        return self._level

    depth = level

    def __repr__(self):
        return "MediaWikiSection(title=%s)" % repr(self.title)


class MediaWikiTable(object):

    def __init__(self, section, title="", headers=[], rows=[], source=""):
        """ A <table class="wikitable> in a MediaWikiSection.
        """
        self.section = section # MediaWikiSection the table is part of.
        self.source  = source  # Table HTML.
        self.title   = title   # Table title.
        self.headers = headers # List of table headers.
        self.rows    = rows    # List of table rows, each a list of cells.

    def plaintext(self, **kwargs):
        return self.article._plaintext(self.source, **kwargs)

    @property
    def html(self):
        return self.source

    @property
    def src(self):
        return self.source

    @property
    def string(self):
        return self.plaintext()

    def __repr__(self):
        return "MediaWikiTable(title=%s)" % repr(self.title)

#--- MEDIAWIKI: WIKIPEDIA --------------------------------------------------------------------------
# Wikipedia is a collaboratively edited, multilingual, free Internet encyclopedia.
# Wikipedia depends on MediaWiki.


class Wikipedia(MediaWiki):

    def __init__(self, license=None, throttle=5.0, language="en"):
        """ Mediawiki search engine for http://[language].wikipedia.org.
        """
        SearchEngine.__init__(self, license or WIKIPEDIA_LICENSE, throttle, language)
        self._subdomain = language

    @property
    def _url(self):
        s = MEDIAWIKI
        s = s.replace("{SUBDOMAIN}", self._subdomain)
        s = s.replace("{DOMAIN}", "wikipedia.org")
        s = s.replace("{API}", "/w/api.php")
        return s

    @property
    def MediaWikiArticle(self):
        return WikipediaArticle

    @property
    def MediaWikiSection(self):
        return WikipediaSection

    @property
    def MediaWikiTable(self):
        return WikipediaTable


class WikipediaArticle(MediaWikiArticle):

    def download(self, media, **kwargs):
        """ Downloads an item from MediaWikiArticle.media and returns the content.
            Note: images on Wikipedia can be quite large, and this method uses screen-scraping,
                  so Wikipedia might not like it that you download media in this way.
            To save the media in a file:
            data = article.download(media)
            open(filename+extension(media),"w").write(data)
        """
        url = "http://%s.wikipedia.org/wiki/File:%s" % (self.__dict__.get("language", "en"), media)
        if url not in cache:
            time.sleep(1)
        data = URL(url).download(**kwargs)
        data = re.search(r"upload.wikimedia.org/.*?/%s" % media, data)
        data = data and URL("http://" + data.group(0)).download(**kwargs) or None
        return data

    def __repr__(self):
        return "WikipediaArticle(title=%s)" % repr(self.title)


class WikipediaSection(MediaWikiSection):
    def __repr__(self):
        return "WikipediaSection(title=%s)" % repr(self.title)


class WikipediaTable(MediaWikiTable):
    def __repr__(self):
        return "WikipediaTable(title=%s)" % repr(self.title)

#article = Wikipedia().search("cat")
#for section in article.sections:
#    print("  "*(section.level-1) + section.title)
#if article.media:
#    data = article.download(article.media[2])
#    f = open(article.media[2], "w")
#    f.write(data)
#    f.close()
#
#article = Wikipedia(language="nl").search("borrelnootje")
#print(article.string)

#for result in Wikipedia().search("\"cat's\"", type="*"):
#    print(result.title)
#    print(result.text)
#    print()

#--- MEDIAWIKI: WIKTIONARY -------------------------------------------------------------------------
# Wiktionary is a collaborative project to produce a free-content multilingual dictionary.


class Wiktionary(MediaWiki):

    def __init__(self, license=None, throttle=5.0, language="en"):
        """ Mediawiki search engine for http://[language].wiktionary.com.
        """
        SearchEngine.__init__(self, license or MEDIAWIKI_LICENSE, throttle, language)
        self._subdomain = language

    @property
    def _url(self):
        s = MEDIAWIKI
        s = s.replace("{SUBDOMAIN}", self._subdomain)
        s = s.replace("{DOMAIN}", "wiktionary.org")
        s = s.replace("{API}", "/w/api.php")
        return s

    @property
    def MediaWikiArticle(self):
        return WiktionaryArticle

    @property
    def MediaWikiSection(self):
        return WiktionarySection

    @property
    def MediaWikiTable(self):
        return WiktionaryTable


class WiktionaryArticle(MediaWikiArticle):
    def __repr__(self):
        return "WiktionaryArticle(title=%s)" % repr(self.title)


class WiktionarySection(MediaWikiSection):
    def __repr__(self):
        return "WiktionarySection(title=%s)" % repr(self.title)


class WiktionaryTable(MediaWikiTable):
    def __repr__(self):
        return "WiktionaryTable(title=%s)" % repr(self.title)


#--- MEDIAWIKI: WIKIA ------------------------------------------------------------------------------
# Wikia (formerly Wikicities) is a free web hosting service and a wiki farm for wikis.
# Wikia hosts several hundred thousand wikis using MediaWiki.

# Author: Robert Elwell (2012)

class Wikia(MediaWiki):

    def __init__(self, domain="www", license=None, throttle=5.0, language="en"):
        """ Mediawiki search engine for http://[domain].wikia.com.
        """
        SearchEngine.__init__(self, license or MEDIAWIKI_LICENSE, throttle, language)
        self._subdomain = domain

    @property
    def _url(self):
        s = MEDIAWIKI
        s = s.replace("{SUBDOMAIN}", self._subdomain)
        s = s.replace("{DOMAIN}", "wikia.com")
        s = s.replace("{API}", '/api.php')
        return s

    @property
    def MediaWikiArticle(self):
        return WikiaArticle

    @property
    def MediaWikiSection(self):
        return WikiaSection

    @property
    def MediaWikiTable(self):
        return WikiaTable

    def articles(self, **kwargs):
        if kwargs.pop("batch", True):
            # We can take advantage of Wikia's search API to reduce bandwith.
            # Instead of executing a query to retrieve each article,
            # we query for a batch of (10) articles.
            iterator = self.index(_id="pageid", **kwargs)
            while True:
                batch, done = [], False
                try:
                    for i in range(10):
                        batch.append(next(iterator))
                except StopIteration:
                    done = True # No more articles, finish batch and raise StopIteration.
                url = URL(self._url.replace("api.php", "wikia.php"), method=GET, query={
                  "controller": "WikiaSearch",
                      "method": "getPages",
                         "ids": '|'.join(str(id) for id in batch),
                      "format": "json"
                })
                kwargs.setdefault("unicode", True)
                kwargs.setdefault("cached", True)
                kwargs["timeout"] = 10 * (1 + len(batch))
                data = url.download(**kwargs)
                data = json.loads(data)
                for x in (data or {}).get("pages", {}).values():
                    yield WikiaArticle(title=x.get("title", ""), source=x.get("html", ""))
                if done:
                    raise StopIteration
        for title in self.index(**kwargs):
            yield self.search(title, **kwargs)


class WikiaArticle(MediaWikiArticle):
    def __repr__(self):
        return "WikiaArticle(title=%s)" % repr(self.title)


class WikiaSection(MediaWikiSection):
    def __repr__(self):
        return "WikiaSection(title=%s)" % repr(self.title)


class WikiaTable(MediaWikiTable):
    def __repr__(self):
        return "WikiaTable(title=%s)" % repr(self.title)

#--- DBPEDIA --------------------------------------------------------------------------------------------------
# DBPedia is a database of structured information mined from Wikipedia.
# DBPedia data is stored as RDF triples: (subject, predicate, object),
# e.g., X is-a Actor, Y is-a Country, Z has-birthplace Country, ...

# DBPedia can be queried using SPARQL:
# http://www.w3.org/TR/rdf-sparql-query/
# A SPARQL query yields rows that match all triples in the WHERE clause.
# A SPARQL query uses ?wildcards in triple subject/object to select fields.

# For example:
# > PREFIX dbo: <http://dbpedia.org/ontology/>
# > SELECT ?actor ?place
# > WHERE {
# >     ?actor a dbo:Actor; dbo:birthPlace ?place.
# >     ?place a dbo:Country.
# > }
#
# - Each row in the results has an "actor" and a "place" field.
# - The actor is of the class "Actor".
# - The place is of the class "Country".
# - Only actors for which a place of birth is known are retrieved.
#
# The fields are RDF resources, e.g.:
# http://dbpedia.org/resource/Australia

# Author: Kenneth Koch (2013) <kkoch986@gmail.com>

DBPEDIA = "http://dbpedia.org/sparql?"

SPARQL = "sparql"


class DBPediaQueryError(HTTP400BadRequest):
    pass


class DBPediaResource(str):
    @property
    def name(self):
        # http://dbpedia.org/resource/Australia => Australia
        s = re.sub("^http://dbpedia.org/resource/", "", self)
        s = s.replace("_", " ")
        s = encode_utf8(s)
        s = decode_url(s)
        s = decode_utf8(s)
        return s


class DBPedia(SearchEngine):

    def __init__(self, license=None, throttle=1.0, language=None):
        SearchEngine.__init__(self, license, throttle, language)

    def search(self, query, type=SPARQL, start=1, count=10, sort=RELEVANCY, size=None, cached=False, **kwargs):
        """ Returns a list of results from DBPedia for the given SPARQL query.
            - type : SPARQL,
            - start: no maximum,
            - count: maximum 1000,
            There is a limit of 10 requests/second.
            Maximum query execution time is 120 seconds.
        """
        if type not in (SPARQL,):
            raise SearchEngineTypeError
        if not query or count < 1 or start < 1:
            return Results(DBPEDIA, query, type)
        # 1) Construct request URL.
        url = URL(DBPEDIA, method=GET)
        url.query = {
            "format": "json",
             "query": "%s OFFSET %s LIMIT %s" % (query,
                        (start - 1) * min(count, 1000),
                        (start - 0) * min(count, 1000)
            )
        }
        # 2) Parse JSON response.
        try:
            data = URL(url).download(cached=cached, timeout=30, **kwargs)
            data = json.loads(data)
        except HTTP400BadRequest as e:
            raise DBPediaQueryError(e.src.read().splitlines()[0])
        except HTTP403Forbidden:
            raise SearchEngineLimitError
        results = Results(DBPEDIA, url.query, type)
        results.total = None
        for x in data["results"]["bindings"]:
            r = Result(url=None)
            for k in data["head"]["vars"]:
                t1 = x[k].get("type", "literal") # uri | literal | typed-literal
                t2 = x[k].get("datatype", "?")   # http://www.w3.org/2001/XMLSchema#float | int | date
                v = x[k].get("value")
                v = self.format(v)
                if t1 == "uri":
                    v = DBPediaResource(v)
                if t2.endswith("float"):
                    v = float(v)
                if t2.endswith("int"):
                    v = int(v)
                dict.__setitem__(r, k, v)
            results.append(r)
        return results

#--- FLICKR ----------------------------------------------------------------------------------------
# Flickr is a popular image hosting and video hosting website.
# http://www.flickr.com/services/api/

FLICKR = "https://api.flickr.com/services/rest/"
FLICKR_LICENSE = api.license["Flickr"]

INTERESTING = "interesting"


class Flickr(SearchEngine):

    def __init__(self, license=None, throttle=5.0, language=None):
        SearchEngine.__init__(self, license or FLICKR_LICENSE, throttle, language)

    def search(self, query, type=IMAGE, start=1, count=10, sort=RELEVANCY, size=None, cached=True, **kwargs):
        """ Returns a list of results from Flickr for the given query.
            Retrieving the URL of a result (i.e. image) requires an additional query.
            - type : SEARCH, IMAGE,
            - start: maximum undefined,
            - count: maximum 500,
            - sort : RELEVANCY, LATEST or INTERESTING.
            There is no daily limit.
        """
        if type not in (SEARCH, IMAGE):
            raise SearchEngineTypeError
        if not query or count < 1 or start < 1 or start > 500 / count:
            return Results(FLICKR, query, IMAGE)
        # 1) Construct request URL.
        url = FLICKR + "?"
        url = URL(url, method=GET, query={
           "api_key": self.license or "",
            "method": "flickr.photos.search",
              "text": query.replace(" ", "_"),
              "page": start,
          "per_page": min(count, 500),
              "sort": {RELEVANCY: "relevance",
                           LATEST: "date-posted-desc",
                      INTERESTING: "interestingness-desc"}.get(sort)
        })
        if kwargs.get("copyright", True) is False:
            # With copyright=False, only returns Public Domain and Creative Commons images.
            # http://www.flickr.com/services/api/flickr.photos.licenses.getInfo.html
            # 5: "Attribution-ShareAlike License"
            # 7: "No known copyright restriction"
            url.query["license"] = "5,7"
        # 2) Parse XML response.
        kwargs.setdefault("unicode", True)
        kwargs.setdefault("throttle", self.throttle)
        data = url.download(cached=cached, **kwargs)
        data = xml.dom.minidom.parseString(bytestring(data))
        results = Results(FLICKR, query, IMAGE)
        results.total = int(data.getElementsByTagName("photos")[0].getAttribute("total"))
        for x in data.getElementsByTagName("photo"):
            r = FlickrResult(url=None)
            r.__dict__["_id"] = x.getAttribute("id")
            r.__dict__["_size"] = size
            r.__dict__["_license"] = self.license
            r.__dict__["_throttle"] = self.throttle
            r.text = self.format(x.getAttribute("title"))
            r.author = self.format(x.getAttribute("owner"))
            results.append(r)
        return results


class FlickrResult(Result):

    @property
    def url(self):
        # Retrieving the url of a FlickrResult (i.e. image location) requires another query.
        # Note: the "Original" size no longer appears in the response,
        # so Flickr might not like it if we download it.
        url = FLICKR + "?method=flickr.photos.getSizes&photo_id=%s&api_key=%s" % (self._id, self._license)
        data = URL(url).download(throttle=self._throttle, unicode=True)
        data = xml.dom.minidom.parseString(bytestring(data))
        size = {TINY: "Thumbnail",
                SMALL: "Small",
               MEDIUM: "Medium",
                LARGE: "Original"}.get(self._size, "Medium")
        for x in data.getElementsByTagName("size"):
            if size == x.getAttribute("label"):
                return x.getAttribute("source")
            if size == "Original":
                url = x.getAttribute("source")
                url = url[:-len(extension(url)) - 2] + "_o" + extension(url)
                return u(url)

#images = Flickr().search("kitten", count=10, size=SMALL)
#for img in images:
#    print(bytestring(img.description))
#    print(img.url)
#
#data = img.download()
#f = open("kitten"+extension(img.url), "wb")
#f.write(data)
#f.close()

#--- FACEBOOK --------------------------------------------------------------------------------------
# Facebook is a popular online social networking service.
# https://developers.facebook.com/docs/reference/api/

FACEBOOK = "https://graph.facebook.com/"
FACEBOOK_LICENSE = api.license["Facebook"]

FEED     = "feed"      # Facebook timeline.
COMMENTS = "comments"  # Facebook comments (for a given news feed post).
LIKES    = "likes"     # Facebook likes (for a given post or comment).
FRIENDS  = "friends"   # Facebook friends (for a given profile id).


class FacebookResult(Result):
    def __repr__(self):
        return "Result(id=%s)" % repr(self.id)


class Facebook(SearchEngine):

    def __init__(self, license=None, throttle=1.0, language=None):
        SearchEngine.__init__(self, license or FACEBOOK_LICENSE, throttle, language)

    @property
    def _token(self):
        # Yields the "application access token" (stored in api.license["Facebook"]).
        # With this license, we can view public content.
        # To view more information, we need a "user access token" as license key.
        # This token can be retrieved manually from:
        #  http://www.clips.ua.ac.be/pattern-facebook
        # Or parsed from this URL:
        #  https://graph.facebook.com/oauth/authorize?type=user_agent
        #   &client_id=332061826907464
        #   &redirect_uri=http://www.clips.ua.ac.be/pattern-facebook
        #   &scope=read_stream,user_birthday,user_likes,user_photos,friends_birthday,friends_likes
        # The token is valid for a limited duration.
        return URL(FACEBOOK + "oauth/access_token?", query={
               "grant_type": "client_credentials",
                "client_id": "332061826907464",
            "client_secret": "81ff4204e73ecafcd87635a3a3683fbe"
        }).download().split("=")[1]

    def search(self, query, type=SEARCH, start=1, count=10, cached=False, **kwargs):
        """ Returns a list of results from Facebook public status updates for the given query.
            - query: string, or Result.id for NEWS and COMMENTS,
            - type : SEARCH,
            - start: 1,
            - count: maximum 100 for SEARCH and NEWS, 1000 for COMMENTS and LIKES.
            There is an hourly limit of +-600 queries (actual amount undisclosed).
        """
        # Facebook.search(type=SEARCH) returns public posts + author.
        # Facebook.search(type=NEWS) returns posts for the given author (id | alias | "me").
        # Facebook.search(type=COMMENTS) returns comments for the given post id.
        # Facebook.search(type=LIKES) returns authors for the given author, post or comments.
        # Facebook.search(type=FRIENDS) returns authors for the given author.
        # An author is a Facebook user or other entity (e.g., a product page).
        if type not in (SEARCH, NEWS, COMMENTS, LIKES, FRIENDS):
            raise SearchEngineTypeError
        if type in (SEARCH, NEWS):
            max = 100
        if type in (COMMENTS, LIKES):
            max = 1000
        if type in (FRIENDS,):
            max = 10000
        if not query or start < 1 or count < 1:
            return Results(FACEBOOK, query, SEARCH)
        if isinstance(query, FacebookResult):
            query = query.id
        # 1) Construct request URL.
        if type == SEARCH:
            url = FACEBOOK + type
            url = URL(url, method=GET, query={
                         "q": query,
                      "type": "post",
              "access_token": self.license,
                    "offset": (start - 1) * min(count, max),
                     "limit": (start - 0) * min(count, max)
            })
        if type in (NEWS, FEED, COMMENTS, LIKES, FRIENDS):
            url = FACEBOOK + (u(query) or "me").replace(FACEBOOK, "") + "/" + type.replace("news", "feed")
            url = URL(url, method=GET, query={
              "access_token": self.license,
                    "offset": (start - 1) * min(count, max),
                     "limit": (start - 0) * min(count, max),
            })
        if type in (SEARCH, NEWS, FEED):
            url.query["fields"] = ",".join((
                "id", "from", "name", "story", "message", "link", "picture", "created_time", "shares",
                "comments.limit(1).summary(true)",
                   "likes.limit(1).summary(true)"
            ))
        # 2) Parse JSON response.
        kwargs.setdefault("cached", cached)
        kwargs.setdefault("unicode", True)
        kwargs.setdefault("throttle", self.throttle)
        try:
            data = URL(url).download(**kwargs)
        except HTTP400BadRequest:
            raise HTTP401Authentication
        data = json.loads(data)
        results = Results(FACEBOOK, query, SEARCH)
        results.total = None
        for x in data.get("data", []):
            r = FacebookResult(url=None)
            r.id       = self.format(x.get("id"))
            r.url      = self.format(x.get("link"))
            r.text     = self.format(x.get("story", x.get("message", x.get("name"))))
            r.date     = self.format(x.get("created_time"))
            r.votes    = self.format(x.get("like_count", x.get("likes", {}).get("summary", {}).get("total_count", 0)) + 0)
            r.shares   = self.format(x.get("shares", {}).get("count", 0))
            r.comments = self.format(x.get("comments", {}).get("summary", {}).get("total_count", 0) + 0)
            r.author   = self.format(x.get("from", {}).get("id", "")), \
                         self.format(x.get("from", {}).get("name", ""))
            # Set Result.text to author name for likes.
            if type in (LIKES, FRIENDS):
                r.author = \
                   self.format(x.get("id", "")), \
                   self.format(x.get("name", ""))
                r.text = self.format(x.get("name"))
            # Set Result.url to full-size image.
            if re.match(r"^http(s?)://www\.facebook\.com/photo", r.url) is not None:
                r.url = x.get("picture", "").replace("_s", "_b") or r.url
            # Set Result.title to object id.
            if re.match(r"^http(s?)://www\.facebook\.com/", r.url) is not None:
                r.title = r.url.split("/")[-1].split("?")[0]
            results.append(r)
        return results

    def profile(self, id=None, **kwargs):
        """ Returns a Result for the given author id or alias.
        """
        # 1) Construct request URL.
        url = FACEBOOK + (u(id or "me")).replace(FACEBOOK, "")
        url = URL(url, method=GET, query={"access_token": self.license})
        kwargs.setdefault("cached", True)
        kwargs.setdefault("unicode", True)
        kwargs.setdefault("throttle", self.throttle)
        # 2) Parse JSON response.
        try:
            data = URL(url).download(**kwargs)
            data = json.loads(data)
        except HTTP400BadRequest:
            raise HTTP401Authentication
        return Result(
                id = data.get("id", ""),                   # 123456...
               url = data.get("link", ""),                 # https://www.facebook.com/tomdesmedt
            handle = data.get("username", ""),             # tomdesmedt
              name = data.get("name"),                     # Tom De Smedt
              text = data.get("description", ""),          # Artist, scientist, software engineer
          language = data.get("locale", "").split("_")[0], # en_US
              date = data.get("birthday", ""),             # 10/10/1000
            gender = data.get("gender", "")[:1],           # m
            locale = data.get("hometown", {}).get("name", ""),
             votes = int(data.get("likes", 0)) # (for product pages)
        )

    page = profile

#--- PRODUCT REVIEWS -------------------------------------------------------------------------------
# ProductWiki is an open web-based product information resource.
# http://connect.productwiki.com/connect-api/

PRODUCTWIKI = "http://api.productwiki.com/connect/api.aspx"
PRODUCTWIKI_LICENSE = api.license["ProductWiki"]


class ProductWiki(SearchEngine):

    def __init__(self, license=None, throttle=5.0, language=None):
        SearchEngine.__init__(self, license or PRODUCTWIKI_LICENSE, throttle, language)

    def search(self, query, type=SEARCH, start=1, count=10, sort=RELEVANCY, size=None, cached=True, **kwargs):
        """ Returns a list of results from Productwiki for the given query.
            Each Result.reviews is a list of (review, score)-items.
            - type : SEARCH,
            - start: maximum undefined,
            - count: 20,
            - sort : RELEVANCY.
            There is no daily limit.
        """
        if type != SEARCH:
            raise SearchEngineTypeError
        if not query or start < 1 or count < 1:
            return Results(PRODUCTWIKI, query, type)
        # 1) Construct request URL.
        url = PRODUCTWIKI + "?"
        url = URL(url, method=GET, query={
               "key": self.license or "",
                 "q": query,
              "page": start,
                "op": "search",
            "fields": "proscons", # "description,proscons" is heavy.
            "format": "json"
        })
        # 2) Parse JSON response.
        kwargs.setdefault("unicode", True)
        kwargs.setdefault("throttle", self.throttle)
        data = URL(url).download(cached=cached, **kwargs)
        data = json.loads(data)
        results = Results(PRODUCTWIKI, query, type)
        results.total = None
        for x in data.get("products", [])[:count]:
            r = Result(url=None)
            r.__dict__["title"] = u(x.get("title"))
            r.__dict__["text"] = u(x.get("text"))
            r.__dict__["reviews"] = []
            reviews = x.get("community_review") or {}
            for p in reviews.get("pros", []):
                r.reviews.append((p.get("text", ""), int(p.get("score")) or +1))
            for p in reviews.get("cons", []):
                r.reviews.append((p.get("text", ""), int(p.get("score")) or -1))
            r.__dict__["score"] = int(sum(score for review, score in r.reviews))
            results.append(r)
        # Highest score first.
        results.sort(key=lambda r: r.score, reverse=True)
        return results

# Backwards compatibility.
Products = ProductWiki

#for r in ProductWiki().search("tablet"):
#    print(r.title)
#    print(r.score)
#    print(r.reviews)
#    print()

#--- NEWS FEED -------------------------------------------------------------------------------------
# Based on the Universal Feed Parser by Mark Pilgrim:
# http://www.feedparser.org/


class Newsfeed(SearchEngine):

    def __init__(self, license=None, throttle=1.0, language=None):
        SearchEngine.__init__(self, license, throttle, language)

    def search(self, query, type=NEWS, start=1, count=10, sort=LATEST, size=SMALL, cached=True, **kwargs):
        """ Returns a list of results from the given RSS or Atom newsfeed URL.
        """
        if type != NEWS:
            raise SearchEngineTypeError
        if not query or start < 1 or count < 1:
            return Results(query, query, NEWS)
        # 1) Construct request URL.
        # 2) Parse RSS/Atom response.
        kwargs.setdefault("unicode", True)
        kwargs.setdefault("throttle", self.throttle)
        tags = kwargs.pop("tags", [])
        data = URL(query).download(cached=cached, **kwargs)
        data = feedparser.parse(data)
        results = Results(query, query, NEWS)
        results.total = None
        for x in data["entries"][:count]:
            s = "\n\n".join([v.get("value") for v in x.get("content", [])]) or x.get("summary")
            r = Result(url=None)
            r.id       = self.format(x.get("id"))
            r.url      = self.format(x.get("link"))
            r.title    = self.format(x.get("title"))
            r.text     = self.format(s)
            r.date     = self.format(x.get("updated"))
            r.author   = self.format(x.get("author"))
            r.language = self.format(x.get("content") and \
                                x.get("content")[0].get("language") or \
                                               data.get("language"))
            for tag in tags:
                # Parse custom tags.
                # Newsfeed.search(tags=["dc:identifier"]) => Result.dc_identifier.
                tag = tag.replace(":", "_")
                r[tag] = self.format(x.get(tag))
            results.append(r)
        return results

feeds = {
          "Nature": "http://feeds.nature.com/nature/rss/current",
         "Science": "http://www.sciencemag.org/rss/podcast.xml",
  "Herald Tribune": "http://www.iht.com/rss/frontpage.xml",
            "TIME": "http://feeds.feedburner.com/time/topstories",
             "CNN": "http://rss.cnn.com/rss/edition.rss",
}

#for r in Newsfeed().search(feeds["Nature"]):
#    print(r.title)
#    print(r.author)
#    print(r.url)
#    print(plaintext(r.text))
#    print()

#--- QUERY -----------------------------------------------------------------------------------------


def query(string, service=GOOGLE, **kwargs):
    """ Returns the list of search query results from the given service.
        For service=WIKIPEDIA, this is a single WikipediaArticle or None.
    """
    service = service.lower()
    if service in (GOOGLE, "google", "g"):
        engine = Google
    if service in (YAHOO, "yahoo", "y!"):
        engine = Yahoo
    if service in (BING, "bing"):
        engine = Bing
    if service in (DUCKDUCKGO, "duckduckgo", "ddg"):
        engine = DuckDuckGo
    if service in (TWITTER, "twitter", "tw"):
        engine = Twitter
    if service in (FACEBOOK, "facebook", "fb"):
        engine = Facebook
    if service in (WIKIPEDIA, "wikipedia", "wp"):
        engine = Wikipedia
    if service in (WIKIA, "wikia"):
        engine = Wikia
    if service in (DBPEDIA, "dbpedia", "dbp"):
        engine = DBPedia
    if service in (FLICKR, "flickr"):
        engine = Flickr
    try:
        kw = {}
        for a in ("license", "throttle", "language"):
            if a in kwargs:
                kw[a] = kwargs.pop(a)
        return engine(kw).search(string, **kwargs)
    except UnboundLocalError:
        raise SearchEngineError("unknown search engine '%s'" % service)

#--- WEB SORT --------------------------------------------------------------------------------------

SERVICES = {
    GOOGLE    : Google,
    YAHOO     : Yahoo,
    FAROO     : Faroo,
    BING      : Bing,
    TWITTER   : Twitter,
    WIKIPEDIA : Wikipedia,
    WIKIA     : Wikia,
    FLICKR    : Flickr,
    FACEBOOK  : Facebook
}


def sort(terms=[], context="", service=GOOGLE, license=None, strict=True, prefix=False, **kwargs):
    """ Returns a list of (percentage, term)-tuples for the given list of terms.
        Sorts the terms in the list according to search result count.
        When a context is defined, sorts according to relevancy to the context, e.g.:
        sort(terms=["black", "green", "red"], context="Darth Vader") =>
        yields "black" as the best candidate, because "black Darth Vader" is more common in search results.
        - terms   : list of search terms,
        - context : term used for sorting,
        - service : web service name (GOOGLE, YAHOO, BING),
        - license : web service license id,
        - strict  : when True the query constructed from term + context is wrapped in quotes.
    """
    service = SERVICES.get(service, SearchEngine)(license, language=kwargs.pop("language", None))
    R = []
    for word in terms:
        q = prefix and (context + " " + word) or (word + " " + context)
        q.strip()
        q = strict and "\"%s\"" % q or q
        t = service in (WIKIPEDIA, WIKIA) and "*" or SEARCH
        r = service.search(q, type=t, count=1, **kwargs)
        R.append(r)
    s = float(sum([r.total or 1 for r in R])) or 1.0
    R = [((r.total or 1) / s, r.query) for r in R]
    R = sorted(R, reverse=kwargs.pop("reverse", True))
    return R

#print(sort(["black", "happy"], "darth vader", GOOGLE))

#### DOCUMENT OBJECT MODEL #########################################################################
# The Document Object Model (DOM) is a cross-platform and language-independent convention
# for representing and interacting with objects in HTML, XHTML and XML documents.
# The pattern.web DOM can be used to traverse HTML source code as a tree of nested elements.
# The pattern.web DOM is based on Beautiful Soup.

# Beautiful Soup is wrapped in DOM, Element and Text classes, resembling the Javascript DOM.
# Beautiful Soup can also be used directly, since it is imported here.
# L. Richardson (2004), http://www.crummy.com/software/BeautifulSoup/

SOUP = (
    BeautifulSoup.BeautifulSoup,
    BeautifulSoup.Tag,
    BeautifulSoup.NavigableString,
    BeautifulSoup.Comment
)

NODE, TEXT, COMMENT, ELEMENT, DOCUMENT = \
    "node", "text", "comment", "element", "document"

#--- NODE ------------------------------------------------------------------------------------------


class Node(object):

    def __init__(self, html, type=NODE, **kwargs):
        """ The base class for Text, Comment and Element.
            All DOM nodes can be navigated in the same way (e.g. Node.parent, Node.children, ...)
        """
        self.type = type
        self._p = not isinstance(html, SOUP) and BeautifulSoup.BeautifulSoup(u(html), "lxml", **kwargs) or html

    @property
    def _beautifulSoup(self):
        # If you must, access the BeautifulSoup object with Node._beautifulSoup.
        return self._p

    def __eq__(self, other):
        # Two Node objects containing the same BeautifulSoup object, are the same.
        return isinstance(other, Node) and hash(self._p) == hash(other._p)

    def _wrap(self, x):
        # Navigating to other nodes yields either Text, Element or None.
        if isinstance(x, BeautifulSoup.Comment):
            return Comment(x)
        if isinstance(x, BeautifulSoup.Declaration):
            return Text(x)
        if isinstance(x, BeautifulSoup.NavigableString):
            return Text(x)
        if isinstance(x, BeautifulSoup.Tag):
            return Element(x)

    @property
    def parent(self):
        return self._wrap(self._p.parent)

    @property
    def children(self):
        return hasattr(self._p, "contents") and [self._wrap(x) for x in self._p.contents] or []

    @property
    def html(self):
        return self.__str__()

    @property
    def source(self):
        return self.__str__()

    @property
    def next_sibling(self):
        return self._wrap(self._p.next_sibling)

    @property
    def previous_sibling(self):
        return self._wrap(self._p.previous_sibling)

    next, prev, previous = \
        next_sibling, previous_sibling, previous_sibling

    def traverse(self, visit=lambda node: None):
        """ Executes the visit function on this node and each of its child nodes.
        """
        visit(self)
        [node.traverse(visit) for node in self.children]

    def remove(self, child):
        """ Removes the given child node (and all nested nodes).
        """
        child._p.extract()

    def __bool__(self):
        return True

    def __len__(self):
        return len(self.children)

    def __iter__(self):
        return iter(self.children)

    def __getitem__(self, index):
        return self.children[index]

    def __repr__(self):
        return "Node(type=%s)" % repr(self.type)

    def __str__(self):
        return u(self._p)

    def __call__(self, *args, **kwargs):
        pass

#--- TEXT ------------------------------------------------------------------------------------------


class Text(Node):
    """ Text represents a chunk of text without formatting in a HTML document.
        For example: "the <b>cat</b>" is parsed to [Text("the"), Element("cat")].
    """

    def __init__(self, string):
        Node.__init__(self, string, type=TEXT)

    def __repr__(self):
        return "Text(%s)" % repr(self._p)


class Comment(Text):
    """ Comment represents a comment in the HTML source code.
        For example: "<!-- comment -->".
    """

    def __init__(self, string):
        Node.__init__(self, string, type=COMMENT)

    def __repr__(self):
        return "Comment(%s)" % repr(self._p)

#--- ELEMENT ---------------------------------------------------------------------------------------


class Element(Node):

    def __init__(self, html):
        """ Element represents an element or tag in the HTML source code.
            For example: "<b>hello</b>" is a "b"-Element containing a child Text("hello").
        """
        Node.__init__(self, html, type=ELEMENT)

    @property
    def tagname(self):
        return self._p.name

    tag = tagName = tagname

    @property
    def attributes(self):
        return self._p.attrs

    attr = attrs = attributes

    @property
    def id(self):
        return self.attributes.get("id")

    @property
    def content(self):
        """ Yields the element content as a unicode string.
        """
        return "".join([u(x) for x in self._p.contents])

    string = content

    @property
    def source(self):
        """ Yields the HTML source as a unicode string (tag + content).
        """
        return u(self._p)

    html = src = source

    def get_elements_by_tagname(self, v):
        """ Returns a list of nested Elements with the given tag name.
            The tag name can include a class (e.g. div.header) or an id (e.g. div#content).
        """
        if isinstance(v, str) and "#" in v:
            v1, v2 = v.split("#")
            v1 = v1 in ("*", "") or v1.lower()
            return [Element(x) for x in self._p.find_all(v1, id=v2)]
        if isinstance(v, str) and "." in v:
            v1, v2 = v.split(".")
            v1 = v1 in ("*", "") or v1.lower()
            return [Element(x) for x in self._p.find_all(v1, v2)]
        return [Element(x) for x in self._p.find_all(v in ("*", "") or v.lower())]

    by_tag = getElementsByTagname = get_elements_by_tagname

    def get_element_by_id(self, v):
        """ Returns the first nested Element with the given id attribute value.
        """
        return ([Element(x) for x in self._p.find_all(id=v, limit=1) or []] + [None])[0]

    by_id = getElementById = get_element_by_id

    def get_elements_by_classname(self, v):
        """ Returns a list of nested Elements with the given class attribute value.
        """
        return [Element(x) for x in (self._p.find_all(True, v))]

    by_class = getElementsByClassname = get_elements_by_classname

    def get_elements_by_attribute(self, **kwargs):
        """ Returns a list of nested Elements with the given attribute value.
        """
        return [Element(x) for x in (self._p.find_all(True, attrs=kwargs))]

    by_attribute = by_attr = getElementsByAttribute = get_elements_by_attribute

    def __call__(self, selector):
        """ Returns a list of nested Elements that match the given CSS selector.
            For example: Element("div#main p.comment a:first-child") matches:
        """
        return SelectorChain(selector).search(self)

    def __getattr__(self, k):
        if k in self.__dict__:
            return self.__dict__[k]
        if k in self.attributes:
            return self.attributes[k]
        raise AttributeError("'Element' object has no attribute '%s'" % k)

    def __contains__(self, v):
        if isinstance(v, Element):
            v = v.content
        return v in self.content

    def __repr__(self):
        return "Element(tag=%s)" % repr(self.tagname)

#--- DOCUMENT --------------------------------------------------------------------------------------


class Document(Element):

    def __init__(self, html, **kwargs):
        """ Document is the top-level element in the Document Object Model.
            It contains nested Element, Text and Comment nodes.
        """
        # Aliases for BeautifulSoup optional parameters:
        # kwargs["selfClosingTags"] = kwargs.pop("self_closing", kwargs.get("selfClosingTags"))
        Node.__init__(self, u(html).strip(), type=DOCUMENT, **kwargs)

    @property
    def declaration(self):
        """ Yields the <!doctype> declaration, as a TEXT Node or None.
        """
        for child in self.children:
            if isinstance(child._p, (BeautifulSoup.Declaration, BeautifulSoup.Doctype)):
                return child

    @property
    def head(self):
        return self._wrap(self._p.head)

    @property
    def body(self):
        return self._wrap(self._p.body)

    @property
    def tagname(self):
        return None

    tag = tagname

    def __repr__(self):
        return "Document()"

DOM = Document

#article = Wikipedia().search("Document Object Model")
#dom = DOM(article.html)
#print(dom.get_element_by_id("References").source)
#print([element.attributes["href"] for element in dom.get_elements_by_tagname("a")])
#print(dom.get_elements_by_tagname("p")[0].next.previous.children[0].parent.__class__)
#print()

#--- DOM CSS SELECTORS -----------------------------------------------------------------------------
# CSS selectors are pattern matching rules (or selectors) to select elements in the DOM.
# CSS selectors may range from simple element tag names to rich contextual patterns.
# http://www.w3.org/TR/CSS2/selector.html

# "*"                 =  <div>, <p>, ...                (all elements)
# "*#x"               =  <div id="x">, <p id="x">, ...  (all elements with id="x")
# "div#x"             =  <div id="x">                   (<div> elements with id="x")
# "div.x"             =  <div class="x">                (<div> elements with class="x")
# "div[class='x']"    =  <div class="x">                (<div> elements with attribute "class"="x")
# "div:contains('x')" =  <div>xyz</div>                 (<div> elements that contain "x")
# "div:first-child"   =  <div><a>1st<a><a></a></div>    (first child inside a <div>)
# "div a"             =  <div><p><a></a></p></div>      (all <a>'s inside a <div>)
# "div, a"            =  <div>, <a>                     (all <a>'s and <div> elements)
# "div + a"           =  <div></div><a></a>             (all <a>'s directly preceded by <div>)
# "div > a"           =  <div><a></a></div>             (all <a>'s directly inside a <div>)
# "div < a"                                            (all <div>'s directly containing an <a>)

# Selectors are case-insensitive.


def _encode_space(s):
    return s.replace(" ", "<!space!>")


def _decode_space(s):
    return s.replace("<!space!>", " ")


class Selector(object):

    def __init__(self, s):
        """ A simple CSS selector is a type (e.g., "p") or universal ("*") selector
            followed by id selectors, attribute selectors, or pseudo-elements.
        """
        self.string = s
        s = s.strip()
        s = s.lower()
        s = s.startswith(("#", ".", ":")) and "*" + s or s
        s = s.replace("#", " #") + " #" # #id
        s = s.replace(".", " .")        # .class
        s = s.replace(":", " :")        # :pseudo-element
        s = s.replace("[", " [")        # [attribute="value"]
        s = re.sub(r"\[.*?\]",
            lambda m: re.sub(r" (\#|\.|\:)", "\\1", m.group(0)), s)
        s = re.sub(r"\[.*?\]",
            lambda m: _encode_space(m.group(0)), s)
        s = re.sub(r":contains\(.*?\)",
            lambda m: _encode_space(m.group(0)), s)
        s = s.split(" ")
        self.tag, self.id, self.classes, self.pseudo, self.attributes = (
             s[0],
              [x[1:] for x in s if x[0] == "#"][0],
          set([x[1:] for x in s if x[0] == "."]),
          set([x[1:] for x in s if x[0] == ":"]),
         dict(self._parse_attribute(x) for x in s if x[0] == "[")
        )

    def _parse_attribute(self, s):
        """ Returns an (attribute, value)-tuple for the given attribute selector.
        """
        s = s.strip("[]")
        s = s.replace("'", "")
        s = s.replace('"', "")
        s = _decode_space(s)
        s = re.sub(r"(\~|\||\^|\$|\*)\=", "=\\1", s)
        s = s.split("=") + [True]
        s = s[:2]
        if s[1] is not True:
            r = r"^%s$"
            if s[1].startswith(("~", "|", "^", "$", "*")):
                p, s[1] = s[1][0], s[1][1:]
                if p == "~":
                    r = r"(^|\s)%s(\s|$)"
                if p == "|":
                    r = r"^%s(-|$)" # XXX doesn't work with spaces.
                if p == "^":
                    r = r"^%s"
                if p == "$":
                    r = r"%s$"
                if p == "*":
                    r = r"%s"
            s[1] = re.compile(r % s[1], re.I)
        return s[:2]

    def _first_child(self, e):
        """ Returns the first child Element of the given element.
        """
        if isinstance(e, Node):
            for e in e.children:
                if isinstance(e, Element):
                    return e

    def _next_sibling(self, e):
        """ Returns the first next sibling Element of the given element.
        """
        while isinstance(e, Node):
            e = e.next_element
            if isinstance(e, Element):
                return e

    def _previous_sibling(self, e):
        """ Returns the last previous sibling Element of the given element.
        """
        while isinstance(e, Node):
            e = e.previous_element
            if isinstance(e, Element):
                return e

    def _contains(self, e, s):
        """ Returns True if string s occurs in the given element (case-insensitive).
        """
        s = re.sub(r"^contains\((.*?)\)$", "\\1", s)
        s = re.sub(r"^[\"']|[\"']$", "", s)
        s = _decode_space(s)
        return re.search(s.lower(), e.content.lower()) is not None

    def match(self, e):
        """ Returns True if the given element matches the simple CSS selector.
        """
        if not isinstance(e, Element):
            return False
        if self.tag not in (e.tag, "*"):
            return False
        if self.id not in ((e.id or "").lower(), "", None):
            return False
        if self.classes.issubset(set(map(lambda s: s.lower(), e.attr.get("class", "")))) is False:
            return False
        if "first-child" in self.pseudo and self._first_child(e.parent) != e:
            return False
        if any(x.startswith("contains") and not self._contains(e, x) for x in self.pseudo):
            return False # jQuery :contains("...") selector.
        for k, v in self.attributes.items():
            if k not in e.attrs or not (v is True or re.search(v, " ".join(e.attrs[k])) is not None):
                return False
        return True

    def search(self, e):
        """ Returns the nested elements that match the simple CSS selector.
        """
        # Map tag to True if it is "*".
        tag = self.tag == "*" or self.tag
        # Map id into a case-insensitive **kwargs dict.
        i = lambda s: re.compile(r"\b%s(?=$|\s)" % s, re.I)
        a = {"id": i(self.id)} if self.id else {}
        a.update(list(map(lambda kv: (kv[0], kv[1]), list(self.attributes.items()))))
        # Match tag + id + all classes + relevant pseudo-elements.
        if not isinstance(e, Element):
            return []
        if len(self.classes) == 0 or len(self.classes) >= 2:
            e = list(map(Element, e._p.find_all(tag, attrs=a)))
        if len(self.classes) == 1:
            e = list(map(Element, e._p.find_all(tag, attrs=dict(a, **{"class": i(list(self.classes)[0])}))))
        if len(self.classes) >= 2:
            e = list(filter(lambda e: self.classes.issubset(set(e.attr.get("class", ""))), e))
        if "first-child" in self.pseudo:
            e = list(filter(lambda e: e == self._first_child(e.parent), e))
        if any(x.startswith("contains") for x in self.pseudo):
            e = list(filter(lambda e: all(not x.startswith("contains") or self._contains(e, x) for x in self.pseudo), e))
        return e

    def __repr__(self):
        return "Selector(%s)" % repr(self.string)


class SelectorChain(list):

    def __init__(self, s):
        """ A selector is a chain of one or more simple selectors,
            separated by combinators (e.g., ">").
        """
        self.string = s
        for s in s.split(","):
            s = s.lower()
            s = s.strip()
            s = re.sub(r" +", " ", s)
            s = re.sub(r" *\> *", " >", s)
            s = re.sub(r" *\< *", " <", s)
            s = re.sub(r" *\+ *", " +", s)
            s = re.sub(r"\[.*?\]",
                lambda m: _encode_space(m.group(0)), s)
            s = re.sub(r":contains\(.*?\)",
                lambda m: _encode_space(m.group(0)), s)
            self.append([])
            for s in s.split(" "):
                if not s.startswith((">", "<", "+")):
                    self[-1].append((" ", Selector(s)))
                elif s.startswith(">"):
                    self[-1].append((">", Selector(s[1:])))
                elif s.startswith("<"):
                    self[-1].append(("<", Selector(s[1:])))
                elif s.startswith("+"):
                    self[-1].append(("+", Selector(s[1:])))

    def search(self, e):
        """ Returns the nested elements that match the CSS selector chain.
        """
        m, root = [], e
        for chain in self:
            e = [root]
            for combinator, s in chain:
                # Search Y, where:
                if combinator == " ":
                    # X Y => X is ancestor of Y
                    e = list(map(s.search, e))
                    e = list(itertools.chain(*e))
                if combinator == ">":
                    # X > Y => X is parent of Y
                    e = list(map(lambda e: list(filter(s.match, e.children)), e))
                    e = list(itertools.chain(*e))
                if combinator == "<":
                    # X < Y => X is child of Y
                    e = list(map(lambda e: e.parent, e))
                    e = list(filter(s.match, e))
                if combinator == "+":
                    # X + Y => X directly precedes Y
                    e = list(map(s._next_sibling, e))
                    e = list(filter(s.match, e))
            m.extend(e)
        return m

#dom = DOM("""
#<html>
#<head></head>
#<body>
#    <div id="#main">
#        <span class="11 22 33">x</span>
#    </div>
#</body>
#</hmtl>
#""")
#
#print(dom("*[class='11']"))
#print(dom("*[class^='11']"))
#print(dom("*[class~='22']"))
#print(dom("*[class$='33']"))
#print(dom("*[class*='3']"))

#### WEB CRAWLER ###################################################################################
# Tested with a crawl across 1,000 domains so far.


class Link(object):

    def __init__(self, url, text="", relation="", referrer=""):
        """ A hyperlink parsed from a HTML document, in the form:
            <a href="url"", title="text", rel="relation">xxx</a>.
        """
        self.url, self.text, self.relation, self.referrer = \
            u(url), u(text), u(relation), u(referrer),

    @property
    def description(self):
        return self.text

    def __repr__(self):
        return "Link(url=%s)" % repr(self.url)

    # Used for sorting in Crawler.links:
    def __eq__(self, link):
        return self.url == link.url

    def __ne__(self, link):
        return self.url != link.url

    def __lt__(self, link):
        return self.url < link.url

    def __gt__(self, link):
        return self.url > link.url


class HTMLLinkParser(HTMLParser):

    def __init__(self):
        HTMLParser.__init__(self)

    def parse(self, html, url=""):
        """ Returns a list of Links parsed from the given HTML string.
        """
        if html is None:
            return None
        self._url = url
        self._data = []
        self.feed(self.clean(html))
        self.close()
        self.reset()
        return self._data

    def handle_starttag(self, tag, attributes):
        if tag == "a":
            attributes = dict(attributes)
            if "href" in attributes:
                link = Link(url = attributes.get("href"),
                           text = attributes.get("title"),
                       relation = attributes.get("rel", ""),
                       referrer = self._url)
                self._data.append(link)


def base(url):
    """ Returns the URL domain name:
        http://en.wikipedia.org/wiki/Web_crawler => en.wikipedia.org
    """
    return urlparse(url).netloc


def abs(url, base=None):
    """ Returns the absolute URL:
        ../media + http://en.wikipedia.org/wiki/ => http://en.wikipedia.org/media
    """
    if url.startswith("#") and base is not None and not base.endswith("/"):
        if not re.search("[^/]/[^/]", base):
            base += "/"
    return urljoin(base, url)

DEPTH = "depth"
BREADTH = "breadth"

FIFO = "fifo" # First In, First Out.
FILO = "filo" # First In, Last Out.
LIFO = "lifo" # Last In, First Out (= FILO).


class Crawler(object):

    def __init__(self, links=[], domains=[], delay=20.0, parse=HTMLLinkParser().parse, sort=FIFO):
        """ A crawler can be used to browse the web in an automated manner.
            It visits the list of starting URLs, parses links from their content, visits those, etc.
            - Links can be prioritized by overriding Crawler.priority().
            - Links can be ignored by overriding Crawler.follow().
            - Each visited link is passed to Crawler.visit(), which can be overridden.
        """
        self.parse    = parse
        self.delay    = delay   # Delay between visits to the same (sub)domain.
        self.domains  = domains # Domains the crawler is allowed to visit.
        self.history  = {}      # Domain name => time last visited.
        self.visited  = {}      # URLs visited.
        self._queue   = []      # URLs scheduled for a visit: (priority, time, Link).
        self._queued  = {}      # URLs scheduled so far, lookup dictionary.
        self.QUEUE    = 10000   # Increase or decrease according to available memory.
        self.sort     = sort
        # Queue given links in given order:
        for link in (isinstance(links, str) and [links] or links):
            self.push(link, priority=1.0, sort=FIFO)

    @property
    def done(self):
        """ Yields True if no further links are scheduled to visit.
        """
        return len(self._queue) == 0

    def push(self, link, priority=1.0, sort=FILO):
        """ Pushes the given link to the queue.
            Position in the queue is determined by priority.
            Equal ranks are sorted FIFO or FILO.
            With priority=1.0 and FILO, the link is inserted to the queue.
            With priority=0.0 and FIFO, the link is appended to the queue.
        """
        if not isinstance(link, Link):
            link = Link(url=link)
        dt = time.time()
        dt = sort == FIFO and dt or 1 / dt
        bisect.insort(self._queue, (1 - priority, dt, link))
        self._queued[link.url] = True

    def pop(self, remove=True):
        """ Returns the next Link queued to visit and removes it from the queue.
            Links on a recently visited (sub)domain are skipped until Crawler.delay has elapsed.
        """
        now = time.time()
        for i, (priority, dt, link) in enumerate(self._queue):
            if self.delay <= now - self.history.get(base(link.url), 0):
                if remove is True:
                    self._queue.pop(i)
                    self._queued.pop(link.url, None)
                return link

    @property
    def next(self):
        """ Returns the next Link queued to visit (without removing it).
        """
        return self.pop(remove=False)

    def crawl(self, method=DEPTH, **kwargs):
        """ Visits the next link in Crawler._queue.
            If the link is on a domain recently visited (< Crawler.delay) it is skipped.
            Parses the content at the link for new links and adds them to the queue,
            according to their Crawler.priority().
            Visited links (and content) are passed to Crawler.visit().
        """
        link = self.pop()
        if link is None:
            return False
        if link.url not in self.visited:
            t = time.time()
            url = URL(link.url)
            if url.mimetype == "text/html":
                try:
                    kwargs.setdefault("unicode", True)
                    html = url.download(**kwargs)
                    for new in self.parse(html, url=link.url):
                        new.url = abs(new.url, base=url.redirect or link.url)
                        new.url = self.normalize(new.url)
                        # 1) Parse new links from HTML web pages.
                        # 2) Schedule unknown links for a visit.
                        # 3) Only links that are not already queued are queued.
                        # 4) Only links for which Crawler.follow() is True are queued.
                        # 5) Only links on Crawler.domains are queued.
                        if new.url == link.url:
                            continue
                        if new.url in self.visited:
                            continue
                        if new.url in self._queued:
                            continue
                        if self.follow(new) is False:
                            continue
                        if self.domains and not base(new.url).endswith(tuple(self.domains)):
                            continue
                        # 6) Limit the queue (remove tail), unless you are Google.
                        if self.QUEUE is not None and \
                           self.QUEUE * 1.25 < len(self._queue):
                            self._queue = self._queue[:self.QUEUE]
                            self._queued.clear()
                            self._queued.update(dict((q[2].url, True) for q in self._queue))
                        # 7) Position in the queue is determined by Crawler.priority().
                        # 8) Equal ranks are sorted FIFO or FILO.
                        self.push(new, priority=self.priority(new, method=method), sort=self.sort)
                    self.visit(link, source=html)
                except URLError:
                    # URL can not be reached (HTTP404NotFound, URLTimeout).
                    self.fail(link)
            else:
                # URL MIME-type is not HTML, don't know how to handle.
                self.fail(link)
            # Log the current time visited for the domain (see Crawler.pop()).
            # Log the URL as visited.
            self.history[base(link.url)] = t
            self.visited[link.url] = True
            return True
        # Nothing happened, we already visited this link.
        return False

    def normalize(self, url):
        """ Called from Crawler.crawl() to normalize URLs.
            For example: return url.split("?")[0]
        """
        # All links pass through here (visited or not).
        # This can be a place to count backlinks.
        return url

    def follow(self, link):
        """ Called from Crawler.crawl() to determine if it should follow this link.
            For example: return "nofollow" not in link.relation
        """
        return True

    def priority(self, link, method=DEPTH):
        """ Called from Crawler.crawl() to determine the priority of this link,
            as a number between 0.0-1.0. Links with higher priority are visited first.
        """
        # Depth-first search dislikes external links to other (sub)domains.
        external = base(link.url) != base(link.referrer)
        if external is True:
            if method == DEPTH:
                return 0.75
            if method == BREADTH:
                return 0.85
        return 0.80

    def visit(self, link, source=None):
        """ Called from Crawler.crawl() when the link is crawled.
            When source=None, the link is not a web page (and was not parsed),
            or possibly a URLTimeout occured (content size too big).
        """
        pass

    def fail(self, link):
        """ Called from Crawler.crawl() for link whose MIME-type could not be determined,
            or which raised a URLError on download.
        """
        pass

Spider = Crawler

#class Polly(Crawler):
#    def visit(self, link, source=None):
#        print("visited:", link.url, "from:", link.referrer)
#    def fail(self, link):
#        print("failed:", link.url)
#
#p = Polly(links=["http://nodebox.net/"], domains=["nodebox.net"], delay=5)
#while not p.done:
#    p.crawl(method=DEPTH, cached=True, throttle=5)

#--- CRAWL FUNCTION --------------------------------------------------------------------------------
# Functional approach to crawling.


def crawl(links=[], domains=[], delay=20.0, parse=HTMLLinkParser().parse, sort=FIFO, method=DEPTH, **kwargs):
    """ Returns a generator that yields (Link, source)-tuples of visited pages.
        When the crawler is idle, it yields (None, None).
    """
    # The scenarios below defines "idle":
    # - crawl(delay=10, throttle=0)
    #   The crawler will wait 10 seconds before visiting the same subdomain.
    #   The crawler will not throttle downloads, so the next link is visited instantly.
    #   So sometimes (None, None) is returned while it waits for an available subdomain.
    # - crawl(delay=0, throttle=10)
    #   The crawler will wait 10 seconds after each and any visit.
    #   The crawler will not delay before visiting the same subdomain.
    #   So usually a result is returned each crawl.next(), but each call takes 10 seconds.
    # - asynchronous(crawl().next)
    #   AsynchronousRequest.value is set to (Link, source) once AsynchronousRequest.done=True.
    #   The program will not halt in the meantime (i.e., the next crawl is threaded).
    crawler = Crawler(links, domains, delay, parse, sort)
    bind(crawler, "visit", \
        lambda crawler, link, source=None: \
            setattr(crawler, "crawled", (link, source))) # Define Crawler.visit() on-the-fly.
    while not crawler.done:
        crawler.crawled = (None, None)
        crawler.crawl(method, **kwargs)
        yield crawler.crawled

#for link, source in crawl("http://www.clips.ua.ac.be/", delay=0, throttle=1, cached=False):
#    print(link)

#g = crawl("http://www.clips.ua.ac.be/"")
#for i in range(10):
#    p = asynchronous(g.next)
#    while not p.done:
#        print("zzz...")
#        time.sleep(0.1)
#    link, source = p.value
#    print(link)


#### DOCUMENT PARSER ###############################################################################
# Not to be confused with Document, which is the top-level element in the HTML DOM.

class DocumentParserError(Exception):
    pass


class DocumentParser(object):

    def __init__(self, path, *args, **kwargs):
        """ Parses a text document (e.g., .pdf or .docx),
            given as a file path or a string.
        """
        self.content = self._parse(path, *args, **kwargs)

    def _open(self, path):
        """ Returns a file-like object with a read() method,
            from the given file path or string.
        """
        if isinstance(path, str) and os.path.exists(path):
            return open(path, "rb")
        if hasattr(path, "read"):
            return path
        return StringIO(path)

    def _parse(self, path, *args, **kwargs):
        """ Returns a plaintext Unicode string parsed from the given document.
        """
        return plaintext(decode_utf8(self.open(path).read()))

    @property
    def string(self):
        return self.content

    def __str__(self):
        return self.content

#--- PDF PARSER ------------------------------------------------------------------------------------
#  Yusuke Shinyama, PDFMiner, http://www.unixuser.org/~euske/python/pdfminer/


class PDFError(DocumentParserError):
    pass


class PDF(DocumentParser):

    def __init__(self, path, output="txt"):
        self.content = self._parse(path, format=output)

    def _parse(self, path, *args, **kwargs):
        # The output is useful for mining but not for display.
        # Alternatively, PDF(format="html") preserves some layout.
        from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
        from pdfminer.pdfpage import PDFPage
        from pdfminer.converter import TextConverter, HTMLConverter
        from pdfminer.layout import LAParams
        try:
            m = PDFResourceManager()
            s = StringIO()
            p = kwargs.get("format", "txt").endswith("html") and HTMLConverter or TextConverter
            p = p(m, s, codec="utf-8", laparams=LAParams())
            interpreter = PDFPageInterpreter(m, p)
            f = self._open(path)
            for page in PDFPage.get_pages(f, maxpages=0, password=""):
                interpreter.process_page(page)
            f.close()
        except Exception as e:
            raise PDFError(str(e))
        s = s.getvalue()
        s = decode_utf8(s)
        s = s.strip()
        s = re.sub(r"([a-z])\-\n", "\\1", s) # Hyphenation.
        s = s.replace("\n\n", "<!-- #p -->") # Paragraphs.
        s = s.replace("\n", " ")
        s = s.replace("<!-- #p -->", "\n\n")
        s = collapse_spaces(s)
        return s

#--- OOXML PARSER ----------------------------------------------------------------------------------
# python-docx, https://github.com/python-openxml/python-docx


class DOCXError(DocumentParserError):
    pass


class DOCX(DocumentParser):

    def _parse(self, path, *args, **kwargs):
        from docx import Document
        document = Document(path)
        try:
            s = [paragraph.text for paragraph in document.paragraphs]
        except Exception as e:
            raise DOCXError(str(e))
        s = "\n\n".join(p for p in s)
        s = decode_utf8(s)
        s = collapse_spaces(s)
        return s

#---------------------------------------------------------------------------------------------------


def parsepdf(path, *args, **kwargs):
    """ Returns the content as a Unicode string from the given .pdf file.
    """
    return PDF(path, *args, **kwargs).content


def parsedocx(path, *args, **kwargs):
    """ Returns the content as a Unicode string from the given .docx file.
    """
    return DOCX(path, *args, **kwargs).content


def parsehtml(path, *args, **kwargs):
    """ Returns the content as a Unicode string from the given .html file.
    """
    return plaintext(DOM(path, *args, **kwargs).body)


def parsedoc(path, format=None):
    """ Returns the content as a Unicode string from the given document (.html., .pdf, .docx).
    """
    if isinstance(path, str):
        if format == "pdf" or path.endswith(".pdf"):
            return parsepdf(path)
        if format == "docx" or path.endswith(".docx"):
            return parsedocx(path)
        if format == "html" or path.endswith((".htm", ".html", ".xhtml")):
            return parsehtml(path)
    # Brute-force approach if the format is unknown.
    for f in (parsepdf, parsedocx, parsehtml):
        try:
            return f(path)
        except:
            pass
